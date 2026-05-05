from __future__ import annotations

from dataclasses import dataclass, field
from io import BufferedIOBase
from pathlib import Path

from docx import Document


@dataclass(frozen=True)
class DocumentNode:
    node_id: str
    node_type: str
    text: str
    parent_id: str | None = None
    scope_id: str | None = None
    path: tuple[str, ...] = ()
    paragraph_indices: tuple[int, ...] = ()
    table_index: int | None = None
    row_index: int | None = None
    cell_index: int | None = None
    style_name: str | None = None
    alignment: str | None = None
    children: tuple[str, ...] = ()


@dataclass
class DocumentGraph:
    nodes: dict[str, DocumentNode] = field(default_factory=dict)

    @classmethod
    def from_docx(cls, path: str | Path | BufferedIOBase) -> "DocumentGraph":
        doc = Document(str(path) if isinstance(path, str | Path) else path)
        graph = cls()
        document_children: list[str] = []

        for idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue

            node_id = f"p:{idx}"
            graph.nodes[node_id] = DocumentNode(
                node_id=node_id,
                node_type="paragraph",
                text=text,
                parent_id="document",
                paragraph_indices=(idx,),
                style_name=paragraph.style.name if paragraph.style is not None else None,
                alignment=str(paragraph.alignment) if paragraph.alignment is not None else None,
            )
            document_children.append(node_id)

        for table_index, table in enumerate(doc.tables):
            table_id = f"t:{table_index}"
            row_ids: list[str] = []
            table_text_parts: list[str] = []

            for row_index, row in enumerate(table.rows):
                row_id = f"t:{table_index}:r:{row_index}"
                cell_ids: list[str] = []
                cell_texts: list[str] = []

                for cell_index, cell in enumerate(row.cells):
                    cell_text = " ".join(
                        paragraph.text.strip()
                        for paragraph in cell.paragraphs
                        if paragraph.text.strip()
                    )
                    cell_id = f"t:{table_index}:r:{row_index}:c:{cell_index}"
                    graph.nodes[cell_id] = DocumentNode(
                        node_id=cell_id,
                        node_type="cell",
                        text=cell_text,
                        parent_id=row_id,
                        table_index=table_index,
                        row_index=row_index,
                        cell_index=cell_index,
                    )
                    cell_ids.append(cell_id)
                    cell_texts.append(cell_text)

                row_text = " | ".join(cell_texts).strip()
                graph.nodes[row_id] = DocumentNode(
                    node_id=row_id,
                    node_type="row",
                    text=row_text,
                    parent_id=table_id,
                    table_index=table_index,
                    row_index=row_index,
                    children=tuple(cell_ids),
                )
                row_ids.append(row_id)
                table_text_parts.append(row_text)

            graph.nodes[table_id] = DocumentNode(
                node_id=table_id,
                node_type="table",
                text="\n".join(table_text_parts),
                parent_id="document",
                table_index=table_index,
                children=tuple(row_ids),
            )
            document_children.append(table_id)

        graph.nodes["document"] = DocumentNode(
            node_id="document",
            node_type="document",
            text="",
            children=tuple(document_children),
        )
        return graph

    def find_by_type(self, node_type: str) -> list[DocumentNode]:
        return [node for node in self.nodes.values() if node.node_type == node_type]

    def find_table_rows_by_ref(self, ref: str) -> list[DocumentNode]:
        needle = ref.strip()
        rows: list[DocumentNode] = []

        for row in self.find_by_type("row"):
            cells = [self.nodes[cell_id] for cell_id in row.children]
            if cells and cells[0].text.strip() == needle:
                rows.append(row)

        return rows
