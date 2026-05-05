from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document

from .base_agent import BaseAgent, normalize_for_match, read_non_empty_paragraphs
from .config import runtime_kwargs
from .schema import AmendmentAnalysis, BaseAnalysis, PipelineValidationReport, ValidationChecklist


class StrictJudgeValidator(BaseAgent):
    def __init__(self, config: dict[str, Any]) -> None:
        super().__init__(name="SDK Validator", **runtime_kwargs(config, "validator"))
        prompts_dir = Path(__file__).parent / "prompts"
        self._system_prompt = (prompts_dir / "judge_system.txt").read_text(encoding="utf-8")
        self._user_template = (prompts_dir / "judge_user_template.txt").read_text(encoding="utf-8")

    def validate(
        self,
        output_doc: Path,
        checklist: ValidationChecklist,
        amendment_analyses: list[AmendmentAnalysis],
        base_analysis: BaseAnalysis,
        operation_statuses: list[str],
        operation_summary: dict[str, int] | None = None,
    ) -> PipelineValidationReport:
        structural_results = self._validate_skeleton_tables(output_doc, checklist)
        structural_ok = all(item["ok"] for item in structural_results)
        paragraph_lines = read_non_empty_paragraphs(output_doc)
        output_lines = _read_lines_with_tables(output_doc)
        output_joined = "\n".join(normalize_for_match(line) for line in output_lines)
        judge_payload = {
            "base_analysis": base_analysis.to_dict(),
            "amendments": [item.to_dict() for item in amendment_analyses],
            "checklist": checklist.to_dict(),
            "operation_statuses": operation_statuses,
            "output_snapshot": output_lines[:250],
        }
        judge_ok = False
        judge_summary = ""
        judge_failures: list[str] = []
        try:
            data, raw = self.call_llm(
                self._system_prompt,
                self._user_template.format_map(
                    {"judge_input_json": json.dumps(judge_payload, ensure_ascii=False, indent=2)}
                ),
                max_tokens=1800,
            )
            judge_ok = bool(data.get("is_valid"))
            judge_summary = str(data.get("summary", ""))
            judge_failures = [str(item) for item in data.get("failures", [])]
        except Exception as exc:
            raw = f"[judge_error] {exc}"
            judge_summary = raw
            judge_failures = [raw]

        total_intents = sum(len(item.intents) for item in amendment_analyses)
        hard_resolution_failure = False
        if total_intents > 0:
            summary = operation_summary or {}
            total_ops = int(summary.get("total", 0))
            resolved_ops = int(summary.get("resolved", 0))
            ambiguous_ops = int(summary.get("ambiguous", 0))
            unsupported_ops = int(summary.get("unsupported", 0))
            all_ambiguous = total_ops > 0 and ambiguous_ops == total_ops
            if (
                total_ops == 0
                or total_ops != total_intents
                or resolved_ops != total_intents
                or ambiguous_ops > 0
                or unsupported_ops > 0
                or all_ambiguous
            ):
                hard_resolution_failure = True
                judge_ok = False
                judge_failures.append(
                    "Hard fail: каждый intent должен иметь resolved operation "
                    f"(intents={total_intents}, total={total_ops}, resolved={resolved_ops}, "
                    f"ambiguous={ambiguous_ops}, unsupported={unsupported_ops})."
                )

        deterministic_ok, deterministic_failures = self._deterministic_intent_checks(
            amendment_analyses,
            output_joined,
            paragraph_lines,
        )
        phrase_coverage_failures = self._check_phrase_replacement_coverage(
            amendment_analyses, operation_statuses
        )
        if phrase_coverage_failures:
            judge_ok = False
            judge_failures.extend(phrase_coverage_failures)
        if not judge_ok and deterministic_ok and not hard_resolution_failure and not phrase_coverage_failures:
            judge_ok = True
            judge_failures.append("LLM-judge fallback: deterministic intent checks passed.")
        elif deterministic_failures:
            judge_failures.extend(deterministic_failures[:10])

        is_valid = (
            structural_ok
            and judge_ok
            and not hard_resolution_failure
        )
        intent_results = []
        for analysis in amendment_analyses:
            for intent in analysis.intents:
                intent_results.append(
                    {
                        "change_id": intent.change_id,
                        "operation_kind": intent.operation_kind,
                        "validated_by_judge": judge_ok,
                    }
                )

        return PipelineValidationReport(
            structural_ok=structural_ok,
            judge_ok=judge_ok,
            is_valid=is_valid,
            skeleton_results=structural_results,
            judge_summary=judge_summary,
            judge_failures=judge_failures,
            intent_results=intent_results,
        )

    def _validate_skeleton_tables(self, output_doc: Path, checklist: ValidationChecklist) -> list[dict[str, Any]]:
        from docx import Document

        document = Document(output_doc)
        table_texts = [
            " ".join(cell.text for row in table.rows for cell in row.cells).lower()
            for table in document.tables
        ]
        results: list[dict[str, Any]] = []
        for check in checklist.checks:
            if check.get("kind") != "service_table_present":
                continue
            details = check.get("details", {})
            expected_documents = [item.lower() for item in details.get("expected_documents", [])]
            found = any(all(label in table_text for label in expected_documents) for table_text in table_texts)
            results.append(
                {
                    "check_id": check["check_id"],
                    "kind": check["kind"],
                    "scope": details.get("scope", ""),
                    "ok": found,
                }
            )
        return results

    def _deterministic_intent_checks(
        self,
        amendment_analyses: list[AmendmentAnalysis],
        output_joined: str,
        paragraph_lines: list[str],
    ) -> tuple[bool, list[str]]:
        failures: list[str] = []
        point_pattern_cache: dict[int, Any] = {}
        for analysis in amendment_analyses:
            for intent in analysis.intents:
                if intent.operation_kind == "append_words_to_point" and intent.appended_words:
                    expected = normalize_for_match(intent.appended_words)
                    if expected and expected not in output_joined:
                        failures.append(f"Intent {intent.change_id}: appended_words not materialized")
                elif intent.operation_kind in {"replace_point", "replace_person_role", "replace_phrase_globally", "repeal_point"}:
                    expected = normalize_for_match(intent.new_text)
                    if expected and expected not in output_joined:
                        failures.append(f"Intent {intent.change_id}: new_text not materialized")
                    if intent.operation_kind == "replace_point" and intent.point_number:
                        import re
                        if intent.point_number not in point_pattern_cache:
                            point_pattern_cache[intent.point_number] = re.compile(rf"^{intent.point_number}\.\s+")
                        pattern = point_pattern_cache[intent.point_number]
                        point_count = sum(1 for line in paragraph_lines if pattern.match(line.strip()))
                        if point_count != 1:
                            failures.append(
                                f"Intent {intent.change_id}: point {intent.point_number} occurrences={point_count} (expected 1)"
                            )
                elif intent.operation_kind == "replace_appendix_block" and intent.new_block_lines:
                    probe_lines = [normalize_for_match(line) for line in intent.new_block_lines if normalize_for_match(line)]
                    if probe_lines:
                        for probe in probe_lines[:3]:
                            if probe not in output_joined:
                                failures.append(f"Intent {intent.change_id}: block line not materialized")
                                break
                elif intent.operation_kind == "append_section_item":
                    expected = normalize_for_match(intent.new_item_text)
                    if expected and expected not in output_joined:
                        failures.append(f"Intent {intent.change_id}: new_item_text not materialized")
        return (len(failures) == 0), failures


    def _check_phrase_replacement_coverage(
        self,
        amendment_analyses: list[AmendmentAnalysis],
        operation_statuses: list[str],
    ) -> list[str]:
        import re
        applied_counts: dict[str, int] = {}
        for status in operation_statuses:
            m = re.match(r"^(\S+):\s+applied:\s+replace_phrase_globally\s+\((\d+)\s+occurrence", status)
            if m:
                applied_counts[m.group(1)] = int(m.group(2))
        failures: list[str] = []
        for analysis in amendment_analyses:
            for intent in analysis.intents:
                if intent.operation_kind != "replace_phrase_globally":
                    continue
                if not intent.old_text:
                    continue
                count = applied_counts.get(intent.change_id)
                if count == 0:
                    phrase = intent.old_text[:60]
                    failures.append(
                        f"Intent {intent.change_id}: replace_phrase_globally "
                        f"— фраза «{phrase}» не найдена в документе (0 вхождений)"
                    )
        return failures


def _read_lines_with_tables(doc_path: Path) -> list[str]:
    lines = read_non_empty_paragraphs(doc_path)
    document = Document(doc_path)
    for table in document.tables:
        for row in table.rows:
            values = [" ".join(cell.text.split()).strip() for cell in row.cells]
            if any(values):
                lines.append(" | ".join(values))
    return lines
