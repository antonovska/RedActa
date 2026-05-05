from __future__ import annotations

from pathlib import Path
import re
from typing import Iterable

from docx import Document

from .base_agent import compact
from .editor_v2 import build_annotation_paragraph
from .schema import ResolvedOperation
from .utils import to_instrumental


class ConsultantMarkerFormatter:
    def format_marker(self, operation: ResolvedOperation) -> str:
        label = to_instrumental(operation.source_document_label)
        if not label:
            return ""

        if operation.operation_kind == "repeal_point":
            return ""

        introduced_label = self._introduced_label(operation.source_document_label)
        if operation.operation_kind == "append_section_item" and operation.paragraph_ordinal is not None:
            return f"(\u0430\u0431\u0437\u0430\u0446 \u0432\u0432\u0435\u0434\u0435\u043d {introduced_label})"

        if operation.operation_kind == "replace_point" and operation.paragraph_ordinal is not None:
            return f"(\u0432 \u0440\u0435\u0434. {label})"

        if operation.subpoint_ref:
            marker = compact(operation.subpoint_ref)
            if operation.operation_kind == "append_section_item":
                return f'(\u043f\u043f. "{marker}" \u0432\u0432\u0435\u0434\u0435\u043d {introduced_label})'
            if operation.operation_kind == "replace_point":
                return f'(\u043f\u043f. "{marker}" \u0432 \u0440\u0435\u0434. {label})'

        return f"(\u0432 \u0440\u0435\u0434. {label})"

    def _introduced_label(self, label: str) -> str:
        value = compact(label)
        if value.startswith("\u041f\u0440\u0438\u043a\u0430\u0437 "):
            return "\u041f\u0440\u0438\u043a\u0430\u0437\u043e\u043c " + value[len("\u041f\u0440\u0438\u043a\u0430\u0437 "):]
        return to_instrumental(value)


class RevisionMarkerInserter:
    def __init__(self, formatter: ConsultantMarkerFormatter | None = None) -> None:
        self._formatter = formatter or ConsultantMarkerFormatter()

    def insert_markers(self, document_path: Path, operations: Iterable[ResolvedOperation]) -> list[dict[str, object]]:
        document = Document(document_path)
        planned: list[tuple[int, ResolvedOperation, str]] = []
        for operation in operations:
            if operation.status != "resolved" or not operation.paragraph_indices:
                continue
            marker = self._formatter.format_marker(operation)
            if not marker:
                continue
            paragraph_index = self._marker_anchor_index(document, operation)
            if not (0 <= paragraph_index < len(document.paragraphs)):
                continue
            planned.append((paragraph_index, operation, marker))

        inserted: list[dict[str, object]] = []
        seen: set[tuple[int, str]] = set()
        for paragraph_index, operation, marker in sorted(planned, key=lambda item: item[0], reverse=True):
            key = (paragraph_index, marker)
            if key in seen:
                continue
            seen.add(key)
            paragraph = document.paragraphs[paragraph_index]
            if self._next_paragraph_text(paragraph) == marker:
                continue
            paragraph._p.addnext(build_annotation_paragraph(marker))
            inserted.append(
                {
                    "operation_id": operation.operation_id,
                    "paragraph_index": paragraph_index,
                    "marker": marker,
                }
            )
        document.save(document_path)
        return list(reversed(inserted))

    def _marker_anchor_index(self, document: Document, operation: ResolvedOperation) -> int:
        paragraph_index = operation.paragraph_indices[0]
        if operation.operation_kind != "replace_phrase_globally":
            return paragraph_index
        return self._structural_block_end_index(document, paragraph_index)

    def _structural_block_end_index(self, document: Document, paragraph_index: int) -> int:
        if not (0 <= paragraph_index < len(document.paragraphs)):
            return paragraph_index
        top_point_pattern = re.compile(r"^\d+\.\s+")
        subpoint_pattern = re.compile(r"^[\u0430-\u044f\u0451a-z](?:\(\d+\))?\)\s+", re.IGNORECASE)
        paragraphs = document.paragraphs
        start_index = paragraph_index
        text = paragraphs[start_index].text.strip()
        if not top_point_pattern.match(text) and not subpoint_pattern.match(text):
            for idx in range(start_index - 1, -1, -1):
                previous_text = paragraphs[idx].text.strip()
                if not previous_text:
                    continue
                if top_point_pattern.match(previous_text) or subpoint_pattern.match(previous_text):
                    start_index = idx
                break

        end_index = paragraph_index
        for idx in range(paragraph_index + 1, len(paragraphs)):
            text = paragraphs[idx].text.strip()
            if not text:
                continue
            if top_point_pattern.match(text) or subpoint_pattern.match(text):
                break
            end_index = idx
        return end_index

    def _next_paragraph_text(self, paragraph) -> str:
        next_element = paragraph._p.getnext()
        if next_element is None:
            return ""
        text_parts = [
            node.text
            for node in next_element.iter()
            if node.tag.endswith("}t") and node.text
        ]
        return compact("".join(text_parts))
