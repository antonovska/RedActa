from __future__ import annotations

import copy
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.table import _Row, Table
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from .base_agent import delete_paragraph
from .schema import ResolvedOperation
from .utils import format_revision_reference, looks_short_list_item, to_instrumental


def replace_paragraph_text_preserving_ooxml(paragraph: Paragraph, text: str) -> None:
    src_runs = paragraph._p.findall(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"
    )
    for run in src_runs:
        paragraph._p.remove(run)

    new_r = OxmlElement("w:r")
    if src_runs:
        src_rpr = src_runs[0].find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr"
        )
        if src_rpr is not None:
            new_r.append(copy.deepcopy(src_rpr))

    new_t = OxmlElement("w:t")
    new_t.text = text
    if text and (text[0] == " " or text[-1] == " "):
        new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_r.append(new_t)
    paragraph._p.append(new_r)


def set_semicolon_ending(paragraph: Paragraph) -> None:
    text = paragraph.text.rstrip()
    if text.endswith(";"):
        return
    if text.endswith("."):
        replace_paragraph_text_preserving_ooxml(paragraph, text[:-1] + ";")
        return
    replace_paragraph_text_preserving_ooxml(paragraph, text + ";")


def normalize_item_text(text: str) -> str:
    value = text.strip()
    value = re.sub(r";{2,}$", ";", value)
    value = re.sub(r'\."\.$', ".", value)
    return value


def _append_text_run(paragraph_element: OxmlElement, text: str) -> None:
    run = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    text_el.text = text
    if text and (text[0] == " " or text[-1] == " "):
        text_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    run.append(text_el)
    paragraph_element.append(run)


def build_annotation_paragraph(text: str, jc: str = "both") -> OxmlElement:
    paragraph = OxmlElement("w:p")
    paragraph_pr = OxmlElement("w:pPr")
    style = OxmlElement("w:pStyle")
    style.set(qn("w:val"), "0")
    paragraph_pr.append(style)
    justification = OxmlElement("w:jc")
    justification.set(qn("w:val"), jc)
    paragraph_pr.append(justification)
    paragraph.append(paragraph_pr)
    _append_text_run(paragraph, text)
    return paragraph


def clone_paragraph_after(anchor: Paragraph, text: str) -> Paragraph:
    new_p = copy.deepcopy(anchor._p)
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    replace_paragraph_text_preserving_ooxml(paragraph, text)
    return paragraph


class IndexDriftTracker:
    """Tracks cumulative paragraph index shifts as sequential edits change document structure.

    Maintains a live map from original paragraph indices to their current positions
    after insertions and deletions. Handles mid-document insertions correctly:
    only paragraphs after the insertion point are shifted.
    """

    def __init__(self, initial_count: int = 0) -> None:
        # _pos_map[i] = current index of original paragraph i, or -1 if deleted
        self._pos_map: list[int] = list(range(initial_count))
        self._events: list[dict] = []

    def adjust(self, original_index: int) -> int:
        """Return the current document index for an original paragraph index.

        If the paragraph was deleted (pos_map == -1), returns the original index
        unchanged. Callers must guard with `0 <= adjusted < len(document.paragraphs)`
        — a deleted paragraph that shifted others may now be out of range.
        """
        if 0 <= original_index < len(self._pos_map):
            mapped = self._pos_map[original_index]
            return mapped if mapped >= 0 else original_index
        # Extrapolate for indices beyond original count (e.g., append at end)
        tail_shift = (self._pos_map[-1] + 1 - len(self._pos_map)) if self._pos_map else 0
        return original_index + tail_shift

    def record_insert(self, op_id: str, after_current_index: int, count: int) -> None:
        """Record that `count` paragraphs were inserted after `after_current_index`."""
        if count <= 0:
            return
        for i in range(len(self._pos_map)):
            if self._pos_map[i] > after_current_index:
                self._pos_map[i] += count
        self._events.append({
            "op_id": op_id,
            "kind": "insert",
            "after": after_current_index,
            "count": count,
        })

    def record_delete(self, op_id: str, start_current_index: int, count: int) -> None:
        """Record that `count` paragraphs were deleted starting at `start_current_index`."""
        if count <= 0:
            return
        end_current_index = start_current_index + count
        for i in range(len(self._pos_map)):
            pos = self._pos_map[i]
            if start_current_index <= pos < end_current_index:
                self._pos_map[i] = -1
            elif pos >= end_current_index:
                self._pos_map[i] -= count
        self._events.append({
            "op_id": op_id,
            "kind": "delete",
            "start": start_current_index,
            "count": count,
        })

    def record_replace_range(
        self, op_id: str, start_current_index: int, deleted_count: int, inserted_count: int
    ) -> None:
        """Record replacement of a range of paragraphs with new ones."""
        net = inserted_count - deleted_count
        if net == 0 and deleted_count == 0:
            return
        end_current_index = start_current_index + deleted_count
        for i in range(len(self._pos_map)):
            pos = self._pos_map[i]
            if start_current_index <= pos < end_current_index:
                self._pos_map[i] = -1
            elif pos >= end_current_index:
                self._pos_map[i] += net
        self._events.append({
            "op_id": op_id,
            "kind": "replace_range",
            "start": start_current_index,
            "deleted_count": deleted_count,
            "inserted_count": inserted_count,
        })

    def events(self) -> list[dict]:
        return list(self._events)


class EditorV2:
    def edit(self, base_doc: Path, output_doc: Path, operations: list[ResolvedOperation]) -> dict[str, Any]:
        document = Document(base_doc)
        statuses: list[str] = []
        applied_operations: list[ResolvedOperation] = []
        drift = IndexDriftTracker(len(document.paragraphs))

        priority = {
            "replace_appendix_block": 0,
            "replace_point": 1,
            "replace_phrase_globally": 1,
            "append_words_to_point": 1,
            "repeal_point": 1,
            "replace_person_role": 1,
            "insert_point": 2,
            "append_section_item": 2,
            "insert_list_entry": 2,
        }

        for operation in sorted(operations, key=lambda item: priority.get(item.operation_kind, 99)):
            status_count = len(statuses)
            event_count = len(drift.events())
            if operation.status != "resolved":
                statuses.append(f"{operation.operation_id}: skipped: {operation.ambiguity_reason or operation.status}")
                continue
            if operation.operation_kind == "insert_point":
                applied = self._apply_insert_point(document, operation, drift)
                if applied:
                    statuses.append(f"{operation.operation_id}: applied: insert_point {operation.point_number}")
                else:
                    statuses.append(f"{operation.operation_id}: skipped: insert_point index out of range after drift adjustment")
            elif operation.operation_kind == "replace_point":
                self._apply_replace_point(document, operation, drift)
                fallback_note = f" [FALLBACK: {operation.ambiguity_reason}]" if operation.ambiguity_reason else ""
                statuses.append(f"{operation.operation_id}: applied: replace_point {operation.point_number}{fallback_note}")
            elif operation.operation_kind == "replace_phrase_globally":
                count, method = self._apply_replace_phrase_globally(document, operation, drift)
                statuses.append(f"{operation.operation_id}: applied: replace_phrase_globally ({count} occurrences, {method})")
            elif operation.operation_kind == "append_section_item":
                applied = self._apply_append_section_item(document, operation, drift)
                if applied:
                    statuses.append(f"{operation.operation_id}: applied: append_section_item")
                else:
                    statuses.append(f"{operation.operation_id}: skipped: append_section_item no applicable anchor")
            elif operation.operation_kind == "append_words_to_point":
                applied = self._apply_append_words_to_point(document, operation, drift)
                if applied:
                    statuses.append(f"{operation.operation_id}: applied: append_words_to_point {operation.point_ref or operation.point_number}")
                else:
                    statuses.append(f"{operation.operation_id}: skipped: append_words_to_point index out of range after drift adjustment")
            elif operation.operation_kind == "repeal_point":
                applied = self._apply_repeal_point(document, operation, drift)
                if applied:
                    statuses.append(f"{operation.operation_id}: applied: repeal_point {operation.point_ref or operation.point_number}")
                else:
                    statuses.append(f"{operation.operation_id}: skipped: repeal_point index out of range after drift adjustment")
            elif operation.operation_kind == "replace_appendix_block":
                applied = self._apply_replace_appendix_block(document, operation, drift)
                if applied:
                    statuses.append(f"{operation.operation_id}: applied: replace_appendix_block {operation.appendix_number}")
                else:
                    statuses.append(f"{operation.operation_id}: skipped: replace_appendix_block index out of range after drift adjustment")
            elif operation.operation_kind == "insert_list_entry":
                applied = self._apply_insert_list_entry(document, operation, drift)
                if applied:
                    statuses.append(f"{operation.operation_id}: applied: insert_list_entry")
                else:
                    statuses.append(f"{operation.operation_id}: skipped: insert_list_entry index out of range after drift adjustment")
            elif operation.operation_kind == "replace_person_role":
                applied = self._apply_replace_person_role(document, operation, drift)
                if applied:
                    statuses.append(f"{operation.operation_id}: applied: replace_person_role")
                else:
                    statuses.append(f"{operation.operation_id}: skipped: replace_person_role index out of range after drift adjustment")
            elif operation.operation_kind == "out_of_scope":
                statuses.append(f"{operation.operation_id}: skipped: out_of_scope (repeal of external act)")
                continue
            else:
                statuses.append(f"{operation.operation_id}: skipped: unsupported operation_kind")
                continue
            self._shift_previous_operation_anchors(applied_operations, drift.events()[event_count:])
            if len(statuses) > status_count and f"{operation.operation_id}: applied:" in statuses[-1]:
                applied_operations.append(operation)

        output_doc.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_doc)
        return {
            "statuses": statuses,
            "drift_events": drift.events(),
            "applied_operations": [self._operation_log_entry(operation) for operation in applied_operations],
        }

    def _shift_previous_operation_anchors(self, operations: list[ResolvedOperation], events: list[dict]) -> None:
        for event in events:
            if event["kind"] == "insert":
                after = event["after"]
                count = event["count"]
                for operation in operations:
                    operation.paragraph_indices = [
                        index + count if index > after else index
                        for index in operation.paragraph_indices
                    ]
            elif event["kind"] == "delete":
                start = event["start"]
                end = start + event["count"]
                for operation in operations:
                    operation.paragraph_indices = [
                        start if start <= index < end else index - event["count"] if index >= end else index
                        for index in operation.paragraph_indices
                    ]
            elif event["kind"] == "replace_range":
                start = event["start"]
                end = start + event["deleted_count"]
                net = event["inserted_count"] - event["deleted_count"]
                for operation in operations:
                    operation.paragraph_indices = [
                        start if start <= index < end else index + net if index >= end else index
                        for index in operation.paragraph_indices
                    ]

    def _operation_log_entry(self, operation: ResolvedOperation) -> dict[str, Any]:
        return {
            "operation_id": operation.operation_id,
            "operation_kind": operation.operation_kind,
            "paragraph_indices": list(operation.paragraph_indices),
            "insert_after_index": operation.insert_after_index,
            "source_document_label": operation.source_document_label,
            "point_ref": operation.point_ref,
            "point_number": operation.point_number,
            "parent_point_ref": operation.parent_point_ref,
            "subpoint_ref": operation.subpoint_ref,
            "paragraph_ordinal": operation.paragraph_ordinal,
            "note_text": operation.note_text,
        }

    def _apply_insert_point(self, document: Document, operation: ResolvedOperation, drift: IndexDriftTracker) -> bool:
        paragraphs = document.paragraphs
        adjusted = drift.adjust(operation.insert_after_index)
        if not (0 <= adjusted < len(paragraphs)):
            return False
        anchor = paragraphs[adjusted]
        if looks_short_list_item(operation.new_text):
            set_semicolon_ending(anchor)
        new_paragraph = clone_paragraph_after(anchor, operation.new_text)
        operation.paragraph_indices = [adjusted + 1]
        drift.record_insert(operation.operation_id, adjusted, 1)
        return True

    def _apply_replace_point(self, document: Document, operation: ResolvedOperation, drift: IndexDriftTracker) -> None:
        point_index = self._resolve_replace_point_index(document, operation, drift)
        paragraph = document.paragraphs[point_index]
        if operation.paragraph_ordinal is not None:
            # Замена конкретного абзаца внутри подпункта ("абзац седьмой подпункта в").
            # Не удаляем другие абзацы, только заменяем текст целевого.
            pass
        elif operation.subpoint_ref:
            # Удаляем только continuation-абзацы самого подпункта (paragraph_indices[1:]),
            # чтобы не затронуть соседние подпункты (г, д, е, ж...).
            # Удаляем в обратном порядке, чтобы не сдвигать индексы.
            deleted_indices: list[int] = []
            for idx in sorted(operation.paragraph_indices[1:], reverse=True):
                adjusted = drift.adjust(idx)
                if 0 <= adjusted < len(document.paragraphs):
                    delete_paragraph(document.paragraphs[adjusted])
                    deleted_indices.append(adjusted)
            if deleted_indices:
                deleted_indices.sort()
                start = deleted_indices[0]
                prev = start
                for curr in deleted_indices[1:] + [None]:
                    if curr is not None and curr == prev + 1:
                        prev = curr
                        continue
                    drift.record_delete(operation.operation_id, start, prev - start + 1)
                    if curr is not None:
                        start = curr
                        prev = curr
        else:  # paragraph_ordinal is guaranteed None here
            deleted_count = self._delete_point_continuation_paragraphs(document, point_index)
            if deleted_count:
                drift.record_delete(operation.operation_id, point_index + 1, deleted_count)
        text = normalize_item_text(operation.new_text)
        if operation.point_number is not None and not operation.subpoint_ref and not text.startswith(f"{operation.point_number}."):
            text = f"{operation.point_number}. {text}"
        if text.count('"') % 2 == 1:
            if text.endswith("."):
                text = text[:-1] + '".'
            else:
                text = text + '"'
        replace_paragraph_text_preserving_ooxml(paragraph, text)
        operation.paragraph_indices = [point_index]

    def _resolve_replace_point_index(self, document: Document, operation: ResolvedOperation, drift: IndexDriftTracker) -> int:
        if operation.point_number and operation.point_number > 0:
            current_idx = self._find_point_block_start_index(document, operation.point_number)
            if current_idx is not None:
                return current_idx
        return drift.adjust(operation.paragraph_indices[0])

    def _apply_replace_phrase_globally(
        self,
        document: Document,
        operation: ResolvedOperation,
        drift: IndexDriftTracker,
    ) -> tuple[int, str]:
        changed_paragraphs: list[tuple[Paragraph, int]] = []

        # Try drift-adjusted indices first.
        for original_index in operation.paragraph_indices:
            adjusted = drift.adjust(original_index)
            if not (0 <= adjusted < len(document.paragraphs)):
                continue
            paragraph = document.paragraphs[adjusted]
            replacement = operation.new_text.upper() if paragraph.text == paragraph.text.upper() else operation.new_text
            updated_text = self._replace_phrase_variants(paragraph.text, operation.old_text, replacement)
            if updated_text is None:
                continue
            replace_paragraph_text_preserving_ooxml(paragraph, updated_text)
            changed_paragraphs.append((paragraph, adjusted))

        method = "adjusted_index"

        # Global fallback: if adjusted indices found nothing, scan the whole document.
        if not changed_paragraphs and operation.old_text:
            method = "global_fallback"
            for current_index, paragraph in enumerate(document.paragraphs):
                replacement = operation.new_text.upper() if paragraph.text == paragraph.text.upper() else operation.new_text
                updated_text = self._replace_phrase_variants(paragraph.text, operation.old_text, replacement)
                if updated_text is None:
                    continue
                replace_paragraph_text_preserving_ooxml(paragraph, updated_text)
                changed_paragraphs.append((paragraph, current_index))

        operation.paragraph_indices = [current_index for _paragraph, current_index in changed_paragraphs]
        return len(changed_paragraphs), method

    def _replace_phrase_variants(self, text: str, old_text: str, replacement: str) -> str | None:
        base_variants: list[str] = [old_text]
        quote_swapped = re.sub(r'"([^"]*)"', r"«\1»", old_text)
        if quote_swapped not in base_variants:
            base_variants.append(quote_swapped)
        ascii_swapped = old_text.replace("«", '"').replace("»", '"')
        if ascii_swapped not in base_variants:
            base_variants.append(ascii_swapped)
        variants: list[str] = []
        for v in base_variants:
            if v not in variants:
                variants.append(v)
            ye = v.replace("ё", "е").replace("Ё", "Е")
            if ye not in variants:
                variants.append(ye)
            if "-" in v:
                for dash in ("–", "—"):
                    dv = v.replace("-", dash)
                    if dv not in variants:
                        variants.append(dv)
            if "–" in v or "—" in v:
                hv = v.replace("–", "-").replace("—", "-")
                if hv not in variants:
                    variants.append(hv)
        norm_text = text.replace("ё", "е").replace("Ё", "Е").lower()
        for variant in variants:
            norm_variant = variant.replace("ё", "е").replace("Ё", "Е").lower()
            if norm_variant not in norm_text:
                continue
            pattern = re.sub(r"[еёЕЁ]", "[её]", re.escape(variant))
            result = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
            if result != text:
                return result
        return None

    def _apply_append_section_item(self, document: Document, operation: ResolvedOperation, drift: IndexDriftTracker) -> bool:
        if operation.new_block_lines:
            applied = self._apply_append_section_item_table(document, operation)
            if applied:
                return True
        anchor_index = self._resolve_insert_anchor_index(document, operation, drift)
        if anchor_index is None:
            return False
        anchor = document.paragraphs[anchor_index]
        if not anchor.text.strip().startswith("("):
            set_semicolon_ending(anchor)
        new_item_text = normalize_item_text(operation.new_item_text)
        split_items = self._split_inline_subitems(new_item_text)
        inserted = 0
        if len(split_items) > 1:
            current_anchor = anchor
            inserted_indices: list[int] = []
            for marker, item_text in split_items:
                new_paragraph = clone_paragraph_after(current_anchor, item_text)
                inserted += 1
                inserted_indices.append(anchor_index + inserted)
                current_anchor = new_paragraph
            drift.record_insert(operation.operation_id, anchor_index, inserted)
            operation.paragraph_indices = inserted_indices
            return True
        new_paragraph = clone_paragraph_after(anchor, new_item_text)
        operation.paragraph_indices = [anchor_index + 1]
        if operation.parent_point_number and operation.subpoint_ref and not self._starts_with_subpoint_ref(new_item_text, operation.subpoint_ref):
            operation.paragraph_ordinal = operation.paragraph_ordinal or 1
        drift.record_insert(operation.operation_id, anchor_index, 1)
        return True

    def _resolve_insert_anchor_index(self, document: Document, operation: ResolvedOperation, drift: IndexDriftTracker) -> int | None:
        if operation.parent_point_number and operation.subpoint_ref:
            if self._starts_with_subpoint_ref(operation.new_item_text, operation.subpoint_ref):
                idx = self._find_previous_subpoint_end_index(document, operation.parent_point_number, operation.subpoint_ref)
            else:
                idx = self._find_subpoint_end_index(document, operation.parent_point_number, operation.subpoint_ref)
            if idx is not None:
                return idx
        if operation.point_number and operation.point_number > 0:
            idx = self._find_point_block_end_index(document, operation.point_number)
            if idx is not None:
                return idx  # live-navigation result — already current, no drift adjustment needed
        adjusted = drift.adjust(operation.insert_after_index)
        if not (0 <= adjusted < len(document.paragraphs)):
            return None
        return adjusted

    def _starts_with_subpoint_ref(self, text: str, subpoint_ref: str) -> bool:
        value = normalize_item_text(text)
        ref = (subpoint_ref or "").strip()
        if not value or not ref:
            return False
        return re.match(rf"^{re.escape(ref)}\)\s+", value, re.IGNORECASE) is not None

    def _find_subpoint_end_index(
        self,
        document: Document,
        parent_point_number: int,
        subpoint_ref: str,
    ) -> int | None:
        top_point_pattern = re.compile(r"^\d+\.\s+")
        subpoint_pattern = re.compile(r"^[а-яёa-z](?:\(\d+\))?\)\s+", re.IGNORECASE)
        target_pattern = re.compile(rf"^{re.escape(subpoint_ref)}\)\s+", re.IGNORECASE)
        point_pattern = re.compile(rf"^{parent_point_number}\.\s+")
        for point_start, paragraph in enumerate(document.paragraphs):
            if not point_pattern.match(paragraph.text.strip()):
                continue
            target_start: int | None = None
            target_end: int | None = None
            for idx in range(point_start + 1, len(document.paragraphs)):
                text = document.paragraphs[idx].text.strip()
                if top_point_pattern.match(text):
                    break
                if target_start is None:
                    if target_pattern.match(text):
                        target_start = idx
                        target_end = idx
                    continue
                if subpoint_pattern.match(text):
                    break
                if text:
                    target_end = idx
            if target_end is not None:
                return target_end
        return None

    def _find_previous_subpoint_end_index(
        self,
        document: Document,
        parent_point_number: int,
        subpoint_ref: str,
    ) -> int | None:
        previous_ref = self._previous_subpoint_ref(subpoint_ref)
        if not previous_ref:
            return None
        top_point_pattern = re.compile(r"^\d+\.\s+")
        subpoint_pattern = re.compile(r"^[а-яёa-z](?:\(\d+\))?\)\s+", re.IGNORECASE)
        previous_pattern = re.compile(rf"^{re.escape(previous_ref)}\)\s+", re.IGNORECASE)
        point_pattern = re.compile(rf"^{parent_point_number}\.\s+")
        for point_start, paragraph in enumerate(document.paragraphs):
            if not point_pattern.match(paragraph.text.strip()):
                continue
            previous_start: int | None = None
            previous_end: int | None = None
            for idx in range(point_start + 1, len(document.paragraphs)):
                text = document.paragraphs[idx].text.strip()
                if top_point_pattern.match(text):
                    break
                if previous_start is None:
                    if previous_pattern.match(text):
                        previous_start = idx
                        previous_end = idx
                    continue
                if subpoint_pattern.match(text):
                    break
                if text:
                    previous_end = idx
            if previous_end is not None:
                return previous_end
        return None

    def _previous_subpoint_ref(self, ref: str) -> str:
        value = (ref or "").strip().lower()
        compound = re.match(r"^([а-яёa-z])\(\d+\)$", value, re.IGNORECASE)
        if compound:
            return compound.group(1).lower()
        letters = "абвгдежзиклмнопрстуфхцчшщэюя"
        if len(value) == 1 and value in letters:
            index = letters.index(value)
            return "" if index == 0 else letters[index - 1]
        return ""

    def _find_point_block_end_index(self, document: Document, point_number: int) -> int | None:
        paragraphs = document.paragraphs
        top_point_pattern = re.compile(r"^(\d+)\.\s+")
        start_idx = None
        for idx, paragraph in enumerate(paragraphs):
            text = paragraph.text.strip()
            match = top_point_pattern.match(text)
            if match and int(match.group(1)) == point_number:
                start_idx = idx
                break
        if start_idx is None:
            return None
        end_idx = start_idx
        for idx in range(start_idx + 1, len(paragraphs)):
            text = paragraphs[idx].text.strip()
            if top_point_pattern.match(text):
                break
            if text:
                end_idx = idx
        return end_idx

    def _find_point_block_start_index(self, document: Document, point_number: int) -> int | None:
        pattern = re.compile(rf"^{point_number}\.\s+")
        for idx, paragraph in enumerate(document.paragraphs):
            if pattern.match(paragraph.text.strip()):
                return idx
        return None

    def _delete_point_continuation_paragraphs(self, document: Document, point_index: int) -> int:
        paragraphs = document.paragraphs
        if point_index < 0 or point_index >= len(paragraphs):
            return 0
        top_point_pattern = re.compile(r"^\d+\.\s+")
        delete_indices: list[int] = []
        for idx in range(point_index + 1, len(paragraphs)):
            text = paragraphs[idx].text.strip()
            if top_point_pattern.match(text):
                break
            if text:
                delete_indices.append(idx)
        for idx in reversed(delete_indices):
            delete_paragraph(document.paragraphs[idx])
        return len(delete_indices)

    def _split_inline_subitems(self, text: str) -> list[tuple[str, str]]:
        value = text.strip()
        if not value:
            return []
        pattern = re.compile(r"(?:(?<=^)|(?<=\s))([а-яёa-z])\)\s*(.+?)(?=(?:\s+[а-яёa-z]\)\s)|$)", re.IGNORECASE | re.DOTALL)
        items: list[tuple[str, str]] = []
        for marker, body in pattern.findall(value):
            item_text = f"{marker}) {' '.join(body.split())}".strip()
            if item_text:
                items.append((marker.lower(), item_text))
        return items

    def _build_subitem_note(self, marker: str, operation: ResolvedOperation) -> str:
        label = to_instrumental(operation.source_document_label)
        if not label:
            return operation.note_text
        return f'(пп. "{marker}" введен {label})'

    def _apply_append_section_item_table(self, document: Document, operation: ResolvedOperation) -> bool:
        table = self._find_target_table_for_append(document, operation)
        if table is None:
            return False

        payload = self._parse_table_payload(operation.new_block_lines)
        if not payload:
            return False

        section_rows = payload["sections"]
        data_rows = payload["rows"]

        section_hint = section_rows[0] if section_rows else ""
        insert_idx = self._find_insert_row_index(table, section_hint, operation.section_hint)
        if insert_idx is None:
            # Fallback: если секция не найдена, ищем строку с номером N-1
            # (предыдущий пункт) по первой ячейке data_rows
            insert_idx = self._find_insert_row_by_prev_number(table, data_rows)
        if insert_idx is None:
            insert_idx = len(table.rows) - 1

        data_template_idx = self._find_data_row_template_index(table, insert_idx)
        last_idx = insert_idx
        for section in section_rows:
            last_idx = self._insert_table_row_after(table, last_idx, [section] * len(table.rows[0].cells))
        for row_values in data_rows:
            padded = list(row_values) + [""] * max(0, len(table.rows[0].cells) - len(row_values))
            last_idx = self._insert_table_row_after(
                table,
                last_idx,
                padded[: len(table.rows[0].cells)],
                template_index=data_template_idx,
            )

        return True

    def _parse_table_payload(self, lines: list[str]) -> dict[str, list[list[str]] | list[str]]:
        sections: list[str] = []
        rows: list[list[str]] = []
        for line in lines:
            if line.startswith("section\t"):
                section = line.split("\t", 1)[1].strip()
                if section:
                    sections.append(section)
                continue
            if line.startswith("row\t"):
                raw_values = line.split("\t")[1:]
                rows.append([value.strip() for value in raw_values])
        return {"sections": sections, "rows": rows}

    def _find_target_table_for_append(self, document: Document, operation: ResolvedOperation) -> Table | None:
        section_hint = operation.section_hint.lower()
        best_table = None
        best_score = -1
        for table in document.tables:
            table_text = " ".join(cell.text for row in table.rows for cell in row.cells).lower()
            score = 0
            if section_hint:
                tokens = [token for token in re.findall(r"[a-zа-я0-9]+", section_hint) if len(token) > 2]
                score = sum(1 for token in tokens if token in table_text)
            if score > best_score:
                best_score = score
                best_table = table
        return best_table

    def _find_insert_row_index(self, table: Table, section_row: str, section_hint: str) -> int | None:
        if not table.rows:
            return None
        normalized_section_row = section_row.lower().strip()
        normalized_hint = section_hint.lower().strip()

        section_idx = None
        if normalized_section_row:
            for idx, row in enumerate(table.rows):
                first = row.cells[0].text.strip().lower()
                if first == normalized_section_row:
                    section_idx = idx
                    break
        if section_idx is None and normalized_hint:
            tokens = [token for token in re.findall(r"[a-zа-я0-9]+", normalized_hint) if len(token) > 2]
            best_idx = None
            best_overlap = 0
            for idx, row in enumerate(table.rows):
                first = row.cells[0].text.strip().lower()
                overlap = sum(1 for token in tokens if token in first)
                if overlap > best_overlap:
                    best_overlap = overlap
                    best_idx = idx
            section_idx = best_idx

        if section_idx is None:
            return None

        insert_idx = section_idx
        for idx in range(section_idx + 1, len(table.rows)):
            if self._is_merged_section_row(table.rows[idx]):
                break
            insert_idx = idx
        return insert_idx

    def _is_merged_section_row(self, row: _Row) -> bool:
        values = [cell.text.strip() for cell in row.cells]
        non_empty = [value for value in values if value]
        if not non_empty:
            return False
        return len(set(non_empty)) == 1

    def _insert_table_row_after(
        self,
        table: Table,
        after_index: int,
        values: list[str],
        template_index: int | None = None,
    ) -> int:
        template_pos = template_index if template_index is not None else after_index
        template_row = table.rows[min(max(template_pos, 0), len(table.rows) - 1)]
        anchor_row = table.rows[min(max(after_index, 0), len(table.rows) - 1)]
        new_tr = copy.deepcopy(template_row._tr)
        anchor_row._tr.addnext(new_tr)
        new_row = _Row(new_tr, table)
        for idx, cell in enumerate(new_row.cells):
            cell.text = values[idx] if idx < len(values) else ""
        return after_index + 1

    def _find_data_row_template_index(self, table: Table, around_index: int) -> int:
        if not table.rows:
            return 0
        for idx in range(around_index + 1, len(table.rows)):
            if not self._is_merged_section_row(table.rows[idx]):
                return idx
        for idx in range(around_index, -1, -1):
            if not self._is_merged_section_row(table.rows[idx]):
                return idx
        return min(max(around_index, 0), len(table.rows) - 1)

    def _find_insert_row_by_prev_number(self, table: Table, data_rows: list[list[str]]) -> int | None:
        """Fallback: ищет строку с номером N-1 для вставки после неё.

        Если первая data_row имеет номер вида '14.' — ищем строку '13.' в таблице
        и вставляем после неё (включая все continuation-строки этого пункта).
        """
        if not data_rows or not data_rows[0]:
            return None
        first_cell = data_rows[0][0].strip()
        m = re.match(r"^(\d+)\.$", first_cell)
        if not m:
            return None
        target_num = int(m.group(1)) - 1
        if target_num < 1:
            return None
        prev_pattern = re.compile(rf"^{target_num}\.$")
        prev_idx = None
        for idx, row in enumerate(table.rows):
            if prev_pattern.match(row.cells[0].text.strip()):
                prev_idx = idx
                break
        if prev_idx is None:
            return None
        # Идём вперёд пока не встретим следующий нумерованный пункт или секцию
        insert_idx = prev_idx
        for idx in range(prev_idx + 1, len(table.rows)):
            first = table.rows[idx].cells[0].text.strip()
            if re.match(r"^\d+\.$", first) or self._is_merged_section_row(table.rows[idx]):
                break
            insert_idx = idx
        return insert_idx

    def _apply_append_words_to_point(self, document: Document, operation: ResolvedOperation, drift: IndexDriftTracker) -> bool:
        adjusted = drift.adjust(operation.paragraph_indices[0])
        if not (0 <= adjusted < len(document.paragraphs)):
            return False
        paragraph = document.paragraphs[adjusted]
        text = paragraph.text.rstrip()
        appended = operation.appended_words.strip()
        if text.endswith("г."):
            text = text[:-2] + " г. " + appended.rstrip(".") + "."
        elif text.endswith(";"):
            text = text[:-1].rstrip() + appended.rstrip(".;") + ";"
        elif text.endswith("."):
            text = text[:-1] + ". " + appended.rstrip(".") + "."
        else:
            text = text + " " + appended.rstrip(".") + "."
        replace_paragraph_text_preserving_ooxml(paragraph, text)
        operation.paragraph_indices = [adjusted]
        return True

    def _is_structural_paragraph(self, paragraph: Paragraph) -> bool:
        import re as _re
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Centered paragraphs (via inline pPr or style inheritance) are always structural:
        # titles, appendix headers, attribution lines, approval stamps are all centered
        # in Russian НПА documents; numbered points are left/justified.
        try:
            if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                return True
            style = paragraph.style
            while style is not None:
                if style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    return True
                style = style.base_style
        except Exception:
            pass

        # Style name indicators (heading/title styles)
        try:
            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            if any(kw in style_name for kw in ("heading", "заголовок", "title", "название")):
                return True
        except Exception:
            pass

        text = paragraph.text.strip()
        if not text:
            return True
        # All-caps lines are structural (document name, section label)
        if text == text.upper() and len(text) > 2 and any(c.isalpha() for c in text):
            return True
        # Appendix header / attribution / approval stamp
        if _re.match(r"^приложение\b", text, _re.IGNORECASE):
            return True
        if _re.match(r"^к\s+(указу|постановлению|приказу|распоряжению|решению)\b", text, _re.IGNORECASE):
            return True
        if _re.match(r"^утвержд", text, _re.IGNORECASE) and len(text) < 100:
            return True
        return False

    def _marker_already_present(self, paragraph: Paragraph, marker_text: str) -> bool:
        next_el = paragraph._p.getnext()
        if next_el is None or next_el.tag != qn("w:p"):
            return False
        p_pr = next_el.find(qn("w:pPr"))
        if p_pr is None:
            return False
        jc = p_pr.find(qn("w:jc"))
        if jc is None or jc.get(qn("w:val")) != "both":
            return False
        text_parts = [
            node.text
            for node in next_el.iter()
            if node.tag == qn("w:t") and node.text
        ]
        return "".join(text_parts) == marker_text

    def _insert_marker_after(
        self,
        paragraph: Paragraph,
        marker_text: str,
        operation: ResolvedOperation,
        drift: IndexDriftTracker,
        current_index: int | None = None,
    ) -> int:
        if not marker_text or self._marker_already_present(paragraph, marker_text):
            return 0
        paragraph._p.addnext(build_annotation_paragraph(marker_text))
        if current_index is not None:
            drift.record_insert(operation.operation_id, current_index, 1)
        return 1

    def _operation_marker_text(self, operation: ResolvedOperation) -> str:
        if operation.note_text:
            return operation.note_text
        label = to_instrumental(operation.source_document_label)
        if not label:
            return ""
        return f"(в ред. {label})"

    def _apply_repeal_point(self, document: Document, operation: ResolvedOperation, drift: IndexDriftTracker) -> bool:
        adjusted = drift.adjust(operation.paragraph_indices[0])
        if not (0 <= adjusted < len(document.paragraphs)):
            return False
        paragraph = document.paragraphs[adjusted]
        replace_paragraph_text_preserving_ooxml(paragraph, operation.new_text)
        operation.paragraph_indices = [adjusted]
        return True

    def _apply_replace_appendix_block(self, document: Document, operation: ResolvedOperation, drift: IndexDriftTracker) -> bool:
        paragraphs = document.paragraphs
        adjusted_start = drift.adjust(operation.paragraph_indices[0])
        if not (0 <= adjusted_start < len(paragraphs)):
            return False
        anchor = paragraphs[adjusted_start - 1] if adjusted_start > 0 else paragraphs[adjusted_start]
        deleted_count = len(paragraphs) - adjusted_start
        for idx in range(len(paragraphs) - 1, adjusted_start - 1, -1):
            delete_paragraph(paragraphs[idx])
        inserted_count = 0
        current_anchor = anchor
        for line in operation.new_block_lines:
            current_anchor = clone_paragraph_after(current_anchor, line)
            inserted_count += 1
        drift.record_replace_range(operation.operation_id, adjusted_start, deleted_count, inserted_count)
        return True

    def _apply_insert_list_entry(self, document: Document, operation: ResolvedOperation, drift: IndexDriftTracker) -> bool:
        adjusted = drift.adjust(operation.insert_after_index)
        if not (0 <= adjusted < len(document.paragraphs)):
            return False
        anchor = document.paragraphs[adjusted]
        text = operation.new_text
        if text and text[-1] not in ".;:":
            text += ";"
        clone_paragraph_after(anchor, text)
        drift.record_insert(operation.operation_id, adjusted, 1)
        return True

    def _apply_replace_person_role(self, document: Document, operation: ResolvedOperation, drift: IndexDriftTracker) -> bool:
        adjusted = drift.adjust(operation.paragraph_indices[0])
        if not (0 <= adjusted < len(document.paragraphs)):
            return False
        paragraph = document.paragraphs[adjusted]
        original = paragraph.text
        if " - " in original:
            prefix, _rest = original.split(" - ", 1)
            new_line = prefix + " - " + operation.new_text
        else:
            new_line = operation.new_text
        if new_line and new_line[-1] not in ".;:":
            new_line += "."
        replace_paragraph_text_preserving_ooxml(paragraph, new_line)
        return True


