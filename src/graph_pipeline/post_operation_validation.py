from __future__ import annotations

from typing import Any

from docx import Document


def _collect_docx_text(path_or_stream: Any) -> str:
    document = Document(path_or_stream)
    parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)

    return "\n".join(parts)


def _materialization_target(operation: Any) -> str:
    new_text = getattr(operation, "new_text", "")
    if new_text:
        return new_text

    new_block_lines = getattr(operation, "new_block_lines", [])
    if new_block_lines:
        return "\n".join(new_block_lines)

    return ""


def validate_operation_materialized(path_or_stream: Any, operation: Any) -> dict[str, Any]:
    status = getattr(operation, "status", "")
    if status != "resolved":
        return {"ok": False, "reason": f"operation status is {status}"}

    expected = _materialization_target(operation)
    if not expected:
        return {"ok": True, "reason": "no materialization target"}

    actual = _collect_docx_text(path_or_stream)
    if expected in actual:
        return {"ok": True, "reason": "materialized"}

    return {"ok": False, "reason": "new_text not materialized"}
