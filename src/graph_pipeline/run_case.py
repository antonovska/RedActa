from __future__ import annotations

import argparse
import json
import re
import time
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document

from .amendment_analyzer import AmendmentAnalyzer
from .base_analyzer import BaseAnalyzer
from .case_loader import load_case
from .config import load_models_config
from .editor import PipelineEditor
from .pipeline_checklist import PipelineChecklist
from .resolver import PipelineResolver
from .service_tables import build_service_table_specs
from .skeleton_builder import SkeletonBuilder
from .validation_checklist_builder import ValidationChecklistBuilder
from .validator import StrictJudgeValidator


def _progress(case_id: str, message: str) -> None:
    ts = datetime.now().strftime("%H:%M:%S")
    print(f"[{ts}][SDK Pipeline][{case_id}] {message}", flush=True)


def _format_seconds(value: float) -> str:
    return f"{value:.1f}с"


def _record_initial_checks(
    checklist: PipelineChecklist,
    *,
    case_topology: str,
    amendment_paths: list[Path],
    base_docs: list[Path],
) -> None:
    checklist.add(
        stage="00_inputs",
        check_id="amendment_documents_present",
        kind="input_documents",
        ok=bool(amendment_paths),
        details={"count": len(amendment_paths), "documents": [path.name for path in amendment_paths]},
    )
    checklist.add(
        stage="00_inputs",
        check_id="base_documents_present",
        kind="input_documents",
        ok=bool(base_docs),
        details={"count": len(base_docs), "documents": [path.name for path in base_docs]},
    )
    checklist.add(
        stage="00_inputs",
        check_id="topology_detected",
        kind="case_topology",
        ok=bool(case_topology),
        details={"case_topology": case_topology},
    )


def _record_analysis_checks(
    checklist: PipelineChecklist,
    *,
    amendment_analyses: list[Any],
    base_analyses: dict[str, Any],
) -> None:
    total_intents = sum(len(item.intents) for item in amendment_analyses)
    checklist.add(
        stage="01_amendment_analysis",
        check_id="amendment_analysis_completed",
        kind="analysis_completed",
        ok=bool(amendment_analyses),
        details={"documents": len(amendment_analyses), "intents": total_intents},
    )
    for analysis in amendment_analyses:
        appendix_targets = sorted({intent.appendix_number for intent in analysis.intents if intent.appendix_number})
        checklist.add(
            stage="01_amendment_analysis",
            check_id=f"amendment_intents_{Path(analysis.metadata.source_path).stem}",
            kind="change_intents_extracted",
            ok=bool(analysis.intents),
            details={
                "source_document_label": analysis.metadata.document_label,
                "target_appendices": appendix_targets,
                "intents": [
                    {
                        "change_id": intent.change_id,
                        "operation_kind": intent.operation_kind,
                        "appendix_number": intent.appendix_number,
                        "point_ref": intent.point_ref,
                        "point_number": intent.point_number,
                        "parent_point_ref": intent.parent_point_ref,
                        "parent_point_number": intent.parent_point_number,
                        "subpoint_ref": intent.subpoint_ref,
                        "section_hint": intent.section_hint,
                        "anchor_text_hint": intent.anchor_text_hint,
                        "source_excerpt": intent.source_excerpt,
                    }
                    for intent in analysis.intents
                ],
            },
        )
        checklist.add(
            stage="01_amendment_analysis",
            check_id=f"target_scopes_{Path(analysis.metadata.source_path).stem}",
            kind="target_scopes_detected",
            ok=bool(analysis.intents),
            details={
                "source_document_label": analysis.metadata.document_label,
                "document_scope_required": True,
                "appendix_scopes_required": appendix_targets,
            },
        )
    checklist.add(
        stage="01_base_analysis",
        check_id="base_analysis_completed",
        kind="analysis_completed",
        ok=bool(base_analyses),
        details={"documents": len(base_analyses)},
    )
    for base_doc, analysis in base_analyses.items():
        checklist.add(
            stage="01_base_analysis",
            check_id=f"base_structure_{Path(base_doc).stem}",
            kind="base_structure_recognized",
            ok=bool(analysis.header_blocks),
            details={
                "base_doc": base_doc,
                "header_blocks": [item.to_dict() for item in analysis.header_blocks],
                "anchors_marked": [
                    {
                        "header_id": item.header_id,
                        "scope": item.scope,
                        "appendix_number": item.appendix_number,
                        "insert_after": item.end_paragraph_index,
                    }
                    for item in analysis.header_blocks
                ],
            },
        )


def _read_doc_lines(doc_path: Path) -> list[str]:
    document = Document(doc_path)
    return [" ".join(paragraph.text.split()) for paragraph in document.paragraphs if " ".join(paragraph.text.split())]


def _estimate_action_units(line: str) -> int:
    lower = line.lower()
    if "исключить" in lower:
        quoted = re.findall(r"[\"«]([^\"»]+)[\"»]", line)
        return max(1, len(quoted))
    if "подпункты" in lower and "изложить" in lower:
        before_action = re.split(r"\bизложить\b", line, flags=re.IGNORECASE)[0]
        quoted = re.findall(r"[\"«]([^\"»]+)[\"»]", before_action)
        return max(1, len(quoted))
    return 1


def _extract_amendment_directives(amendment_doc: Path) -> list[str]:
    lines = _read_doc_lines(amendment_doc)
    action_pattern = re.compile(
        r"исключить|заменить|изложить|дополнить|признать\s+утративш|утратившим\s+силу",
        flags=re.IGNORECASE,
    )
    directives: list[str] = []
    in_changes_block = False
    for line in lines:
        if not in_changes_block:
            if re.match(r"^\d+\.\s+", line):
                in_changes_block = True
            else:
                continue
        if not action_pattern.search(line):
            continue
        units = _estimate_action_units(line)
        if units == 1:
            directives.append(line)
            continue
        for unit_index in range(1, units + 1):
            directives.append(f"{line} [atomic_unit={unit_index}/{units}]")
    return directives


def _analysis_gate_details(
    amendment_analyses: list[Any],
    base_analyses: dict[str, Any],
    amendment_paths: list[Path],
) -> tuple[bool, dict[str, Any]]:
    amendment_docs_without_intents = [
        analysis.metadata.document_label
        for analysis in amendment_analyses
        if not analysis.intents
    ]
    bases_without_structure = [
        base_doc
        for base_doc, analysis in base_analyses.items()
        if not analysis.header_blocks
    ]
    coverage = []
    coverage_failed = []
    analysis_by_path = {str(Path(analysis.metadata.source_path)): analysis for analysis in amendment_analyses}
    for amendment_path in amendment_paths:
        analysis = analysis_by_path.get(str(amendment_path))
        directives = _extract_amendment_directives(amendment_path)
        intent_count = len(analysis.intents) if analysis else 0
        expected_count = len(directives)
        ok = expected_count == 0 or intent_count >= expected_count
        item = {
            "source_path": str(amendment_path),
            "document_label": analysis.metadata.document_label if analysis else amendment_path.name,
            "estimated_directives": expected_count,
            "intent_count": intent_count,
            "ok": ok,
            "directives": directives,
        }
        coverage.append(item)
        if not ok:
            coverage_failed.append(item)
    ok = (
        bool(amendment_analyses)
        and bool(base_analyses)
        and not amendment_docs_without_intents
        and not bases_without_structure
        and not coverage_failed
    )
    return ok, {
        "amendment_docs": len(amendment_analyses),
        "base_docs": len(base_analyses),
        "amendment_docs_without_intents": amendment_docs_without_intents,
        "bases_without_structure": bases_without_structure,
        "coverage": coverage,
        "coverage_failed": coverage_failed,
    }


def _enforce_resolution_gate(
    checklist: PipelineChecklist,
    *,
    case_id: str,
    base_doc: Path,
    amendment_analysis: Any,
    amendment_index: int,
    resolved_operations: list[Any],
    pass_name: str = "initial",
    raise_on_fail: bool = True,
) -> bool:
    expected_ids = [intent.change_id for intent in amendment_analysis.intents]
    operation_ids = [operation.operation_id for operation in resolved_operations]
    missing_ids = [change_id for change_id in expected_ids if change_id not in operation_ids]
    extra_ids = [operation_id for operation_id in operation_ids if operation_id not in expected_ids]
    non_resolved = [
        {
            "operation_id": operation.operation_id,
            "status": operation.status,
            "reason": operation.ambiguity_reason,
        }
        for operation in resolved_operations
        if operation.status != "resolved"
    ]
    ok = not missing_ids and not extra_ids and not non_resolved and len(operation_ids) == len(expected_ids)
    details = {
        "base_doc": str(base_doc),
        "amendment": amendment_analysis.metadata.document_label,
        "expected_intent_ids": expected_ids,
        "operation_ids": operation_ids,
        "missing_intent_ids": missing_ids,
        "extra_operation_ids": extra_ids,
        "non_resolved_operations": non_resolved,
    }
    checklist.add(
        stage="03_resolution",
        check_id=f"resolution_gate_{base_doc.stem}_{amendment_index}_{pass_name}",
        kind="resolution_gate",
        ok=ok,
        details=details,
    )
    if ok:
        return True
    if not raise_on_fail:
        return False
    raise RuntimeError(
        "Resolution gate failed: editor is blocked until every extracted intent has exactly one resolved operation. "
        f"case_id={case_id}, base_doc={base_doc.name}, amendment={amendment_analysis.metadata.document_label}, "
        f"expected={len(expected_ids)}, operations={len(operation_ids)}, "
        f"missing={missing_ids}, non_resolved={non_resolved}"
    )


def _enforce_analysis_gate(
    checklist: PipelineChecklist,
    *,
    case_id: str,
    amendment_analyses: list[Any],
    base_analyses: dict[str, Any],
    amendment_paths: list[Path],
    pass_name: str = "initial",
) -> None:
    ok, details = _analysis_gate_details(amendment_analyses, base_analyses, amendment_paths)
    checklist.add(
        stage="01_analysis_gate",
        check_id=f"analysis_gate_{pass_name}",
        kind="analysis_gate",
        ok=ok,
        details=details,
    )
    if ok:
        return
    raise RuntimeError(
        "Analysis gate failed: pipeline is blocked until amendment intents and base structure are recognized. "
        f"case_id={case_id}, details={details}"
    )


def _repair_analysis_once(
    checklist: PipelineChecklist,
    *,
    case_id: str,
    amendment_analyzer: AmendmentAnalyzer,
    amendment_paths: list[Path],
    amendment_analyses: list[Any],
    base_analyses: dict[str, Any],
) -> list[Any]:
    ok, details = _analysis_gate_details(amendment_analyses, base_analyses, amendment_paths)
    checklist.add(
        stage="01_analysis_gate",
        check_id="analysis_gate_initial",
        kind="analysis_gate",
        ok=ok,
        details=details,
    )
    if ok:
        return amendment_analyses

    failed_by_path = {item["source_path"]: item for item in details["coverage_failed"]}
    repaired: list[Any] = []
    for analysis in amendment_analyses:
        source_path = Path(analysis.metadata.source_path)
        failed_item = failed_by_path.get(str(source_path))
        if not failed_item:
            repaired.append(analysis)
            continue
        print(
            f"[{datetime.now().strftime('%H:%M:%S')}][SDK Pipeline][{case_id}] "
            f"analysis repair: {source_path.name} "
            f"(intents={len(analysis.intents)}, estimated_directives={failed_item['estimated_directives']})",
            flush=True,
        )
        repaired_analysis = amendment_analyzer.repair_analysis(source_path, analysis, failed_item["directives"])
        checklist.add(
            stage="01_analysis_gate",
            check_id=f"analysis_repair_{source_path.stem}",
            kind="analysis_repair_attempted",
            ok=True,
            details={
                "source_path": str(source_path),
                "previous_intents": len(analysis.intents),
                "repaired_intents": len(repaired_analysis.intents),
                "estimated_directives": failed_item["estimated_directives"],
            },
        )
        repaired.append(repaired_analysis)
    _enforce_analysis_gate(
        checklist,
        case_id=case_id,
        amendment_analyses=repaired,
        base_analyses=base_analyses,
        amendment_paths=amendment_paths,
        pass_name="repair",
    )
    return repaired


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Run graph_pipeline on one local case")
    parser.add_argument("--case-id", required=True)
    parser.add_argument("--workspace-root", type=Path, default=Path(__file__).parent)
    parser.add_argument("--models-config", type=Path, default=None)
    parser.add_argument("--output-json", type=Path, default=None)
    return parser.parse_args()


def _run_single_base_flow(
    *,
    case_id: str,
    workspace_root: Path,
    artifacts_dir: Path,
    amendment_analyses: list[Any],
    base_doc: Path,
    base_analysis: Any,
    runtime_checklist: PipelineChecklist,
    base_analyzer: BaseAnalyzer,
    skeleton_builder: SkeletonBuilder,
    resolver: PipelineResolver,
    editor: PipelineEditor,
    checklist_builder: ValidationChecklistBuilder,
    validator: StrictJudgeValidator,
) -> dict[str, Any]:
    artifacts_dir.mkdir(parents=True, exist_ok=True)
    _progress(case_id, "--- Шаг 2/4: Каркас и правки ---")

    # Логируем complexity базового документа и всех amendments
    base_complexity = getattr(base_analysis, "complexity", "plain")
    amendment_complexities = [
        getattr(a.metadata, "complexity", "plain") for a in amendment_analyses
    ]
    effective_complexity = (
        "media_heavy" if "media_heavy" in [base_complexity] + amendment_complexities
        else "table_heavy" if "table_heavy" in [base_complexity] + amendment_complexities
        else "plain"
    )
    _progress(case_id, f"  complexity: base={base_complexity}, amendments={amendment_complexities} → effective={effective_complexity}")
    if effective_complexity == "media_heavy":
        _progress(case_id, "  [WARNING] media_heavy complexity detected — embedded images/drawings require manual review (L3-03)")
    elif effective_complexity == "table_heavy":
        _progress(case_id, "  [INFO] table_heavy complexity — using table-aware resolver/editor path")
    step_build_start = time.perf_counter()
    _progress(case_id, f"  build skeleton -> {artifacts_dir}")
    skeleton_doc = artifacts_dir / "working_skeleton.docx"
    service_table_specs = skeleton_builder.build(base_doc, base_analysis, amendment_analyses, skeleton_doc)
    _progress(case_id, f"  service tables: {len(service_table_specs)}")
    runtime_checklist.add(
        stage="02_skeleton",
        check_id=f"skeleton_created_{base_doc.stem}",
        kind="skeleton_created",
        ok=skeleton_doc.exists(),
        details={"base_doc": str(base_doc), "skeleton_doc": str(skeleton_doc)},
    )
    runtime_checklist.add(
        stage="02_skeleton",
        check_id=f"service_tables_planned_{base_doc.stem}",
        kind="service_tables_planned",
        ok=bool(service_table_specs),
        details={"service_table_specs": [item.to_dict() for item in service_table_specs]},
    )
    runtime_checklist.add(
        stage="02_skeleton",
        check_id=f"document_service_table_planned_{base_doc.stem}",
        kind="document_service_table_planned",
        ok=any(item.scope == "document" for item in service_table_specs),
        details={"document_scope_required": True},
    )
    for spec in service_table_specs:
        if spec.scope != "appendix":
            continue
        runtime_checklist.add(
            stage="02_skeleton",
            check_id=f"appendix_service_table_planned_{base_doc.stem}_{spec.appendix_number}",
            kind="appendix_service_table_planned",
            ok=True,
            details=spec.to_dict(),
        )

    current_doc = skeleton_doc
    steps: list[dict[str, Any]] = []
    all_operations = []
    all_statuses: list[str] = []
    for index, amendment_analysis in enumerate(amendment_analyses, 1):
        _progress(
            case_id,
            f"  amendment {index}/{len(amendment_analyses)}: {amendment_analysis.metadata.document_label}",
        )
        step_output = artifacts_dir / f"working_step_{index}.docx"
        resolution = resolver.resolve(current_doc, amendment_analysis.intents)
        resolved_operations = resolution["resolved_operations"]
        resolved_count = sum(1 for item in resolved_operations if item.status == "resolved")
        ambiguous_count = sum(1 for item in resolved_operations if item.status == "ambiguous")
        unsupported_count = sum(1 for item in resolved_operations if item.status == "unsupported")
        _progress(
            case_id,
            "    "
            f"intents={len(amendment_analysis.intents)}, "
            f"operations={len(resolved_operations)}, "
            f"resolved={resolved_count}, "
            f"ambiguous={ambiguous_count}, "
            f"unsupported={unsupported_count}",
        )
        runtime_checklist.add(
            stage="03_resolution",
            check_id=f"resolution_{base_doc.stem}_{index}",
            kind="resolution_completed",
            ok=bool(resolved_operations),
            details={
                "amendment": amendment_analysis.metadata.document_label,
                "operations": [item.to_dict() for item in resolved_operations],
                "debug_candidates": resolution["debug_candidates"],
                "summary": {
                    "total": len(resolved_operations),
                    "resolved": resolved_count,
                    "ambiguous": ambiguous_count,
                    "unsupported": unsupported_count,
                },
            },
        )
        gate_ok = _enforce_resolution_gate(
            runtime_checklist,
            case_id=case_id,
            base_doc=base_doc,
            amendment_analysis=amendment_analysis,
            amendment_index=index,
            resolved_operations=resolved_operations,
            pass_name="initial",
            raise_on_fail=False,
        )
        if not gate_ok:
            _progress(case_id, "    resolution repair: retry without relevance filter")
            resolution = resolver.resolve(current_doc, amendment_analysis.intents, repair=True)
            resolved_operations = resolution["resolved_operations"]
            resolved_count = sum(1 for item in resolved_operations if item.status == "resolved")
            ambiguous_count = sum(1 for item in resolved_operations if item.status == "ambiguous")
            unsupported_count = sum(1 for item in resolved_operations if item.status == "unsupported")
            _progress(
                case_id,
                "    "
                f"repair_operations={len(resolved_operations)}, "
                f"resolved={resolved_count}, "
                f"ambiguous={ambiguous_count}, "
                f"unsupported={unsupported_count}",
            )
            runtime_checklist.add(
                stage="03_resolution",
                check_id=f"resolution_repair_{base_doc.stem}_{index}",
                kind="resolution_repair_attempted",
                ok=True,
                details={
                    "amendment": amendment_analysis.metadata.document_label,
                    "operations": [item.to_dict() for item in resolved_operations],
                    "debug_candidates": resolution["debug_candidates"],
                    "summary": {
                        "total": len(resolved_operations),
                        "resolved": resolved_count,
                        "ambiguous": ambiguous_count,
                        "unsupported": unsupported_count,
                    },
                },
            )
            _enforce_resolution_gate(
                runtime_checklist,
                case_id=case_id,
                base_doc=base_doc,
                amendment_analysis=amendment_analysis,
                amendment_index=index,
                resolved_operations=resolved_operations,
                pass_name="repair",
            )
        edit_result = editor.edit(current_doc, step_output, resolution["resolved_operations"])
        _progress(
            case_id,
            f"    editor statuses={len(edit_result['statuses'])}, output={step_output.name}",
        )
        runtime_checklist.add(
            stage="04_edit",
            check_id=f"edit_{base_doc.stem}_{index}",
            kind="edit_applied",
            ok=step_output.exists(),
            details={
                "amendment": amendment_analysis.metadata.document_label,
                "output_doc": str(step_output),
                "statuses": list(edit_result["statuses"]),
            },
        )
        steps.append(
            {
                "amendment": amendment_analysis.to_dict(),
                "resolution": {
                    "resolved_operations": [item.to_dict() for item in resolved_operations],
                    "debug_candidates": resolution["debug_candidates"],
                },
                "edit": edit_result,
                "output_doc": str(step_output),
            }
        )
        all_operations.extend(resolved_operations)
        all_statuses.extend(edit_result["statuses"])
        current_doc = step_output

    _progress(case_id, f"  edit steps: {len(steps)}")
    _progress(case_id, "  reanalyze final document")
    final_base_analysis = base_analyzer.analyze(current_doc)
    final_document = Document(current_doc)
    final_specs = build_service_table_specs(final_base_analysis, amendment_analyses)
    final_document.save(current_doc)
    runtime_checklist.add(
        stage="04_edit",
        check_id=f"final_service_tables_{base_doc.stem}",
        kind="final_service_tables_inserted",
        ok=bool(final_specs),
        details={"service_table_specs": [item.to_dict() for item in final_specs], "output_doc": str(current_doc)},
    )
    _progress(case_id, f"  Время: {_format_seconds(time.perf_counter() - step_build_start)}")

    _progress(case_id, "--- Шаг 3/4: Checklist ---")
    step_checklist_start = time.perf_counter()
    _progress(case_id, "build validation checklist")
    checklist = checklist_builder.build(
        base_analysis=final_base_analysis,
        amendment_analyses=amendment_analyses,
        service_table_specs=final_specs,
        resolved_operations=all_operations,
        runtime_checks=runtime_checklist.items(),
    )
    _progress(
        case_id,
        f"  checks: {len(checklist.checks)}",
    )
    _progress(case_id, f"  Время: {_format_seconds(time.perf_counter() - step_checklist_start)}")

    _progress(case_id, "--- Шаг 4/4: Валидатор ---")
    step_validate_start = time.perf_counter()
    _progress(case_id, "run validator")
    operation_summary = {
        "total": len(all_operations),
        "resolved": sum(1 for item in all_operations if item.status == "resolved"),
        "ambiguous": sum(1 for item in all_operations if item.status == "ambiguous"),
        "unsupported": sum(1 for item in all_operations if item.status == "unsupported"),
    }
    _progress(
        case_id,
        "  operation summary: "
        f"total={operation_summary['total']}, "
        f"resolved={operation_summary['resolved']}, "
        f"ambiguous={operation_summary['ambiguous']}, "
        f"unsupported={operation_summary['unsupported']}",
    )
    validation = validator.validate(
        output_doc=current_doc,
        checklist=checklist,
        amendment_analyses=amendment_analyses,
        base_analysis=final_base_analysis,
        operation_statuses=all_statuses,
        operation_summary=operation_summary,
    )
    _progress(
        case_id,
        f"  Валидация: {'OK' if validation.is_valid else 'NOT OK'} "
        f"(structural_ok={validation.structural_ok}, judge_ok={validation.judge_ok})",
    )
    _progress(case_id, f"  Время: {_format_seconds(time.perf_counter() - step_validate_start)}")
    return {
        "case_id": str(case_id),
        "base_doc": str(base_doc),
        "base_analysis": final_base_analysis.to_dict(),
        "service_table_specs": [item.to_dict() for item in final_specs],
        "validation_checklist": checklist.to_dict(),
        "steps": steps,
        "output_doc": str(current_doc),
        "validation": validation.to_dict(),
    }


def _analyze_amendments(amendment_analyzer: AmendmentAnalyzer, amendment_paths: list[Path]) -> list[Any]:
    return amendment_analyzer.analyze_many(amendment_paths)


def _analyze_bases(base_analyzer: BaseAnalyzer, base_docs: list[Path]) -> dict[str, Any]:
    return {str(base_doc): base_analyzer.analyze(base_doc) for base_doc in base_docs}


def run_case(case: dict[str, Any], workspace_root: Path, models_config: Path | None = None) -> dict[str, Any]:
    case_id = str(case["case_id"])
    total_start = time.perf_counter()
    config = load_models_config(models_config)
    amendment_analyzer = AmendmentAnalyzer(config)
    base_analyzer = BaseAnalyzer()
    skeleton_builder = SkeletonBuilder()
    resolver = PipelineResolver(config)
    editor = PipelineEditor()
    checklist_builder = ValidationChecklistBuilder()
    validator = StrictJudgeValidator(config)

    case_topology = case.get("case_topology", "standard_single")
    amendment_paths = case["amendment_docs"]
    artifacts_root = workspace_root / "artifacts" / case_id
    base_docs_for_analysis = case.get("base_docs") or ([case["base_doc"]] if case.get("base_doc") else [])
    runtime_checklist = PipelineChecklist(case_id)
    _record_initial_checks(
        runtime_checklist,
        case_topology=case_topology,
        amendment_paths=amendment_paths,
        base_docs=base_docs_for_analysis,
    )

    _progress(case_id, "=" * 60)
    _progress(case_id, f"Запуск SDK pipeline")
    _progress(case_id, f"  topology: {case_topology}")
    _progress(case_id, f"  amendments: {len(amendment_paths)}")
    _progress(case_id, f"  bases: {len(base_docs_for_analysis)}")
    _progress(case_id, f"  artifacts: {artifacts_root}")
    _progress(case_id, "=" * 60)
    _progress(case_id, "--- Шаг 1/4: Аналитики ---")
    step_analysis_start = time.perf_counter()
    with ThreadPoolExecutor(max_workers=2) as executor:
        amendments_future = executor.submit(_analyze_amendments, amendment_analyzer, amendment_paths)
        bases_future = executor.submit(_analyze_bases, base_analyzer, base_docs_for_analysis)
        amendment_analyses = amendments_future.result()
        base_analyses = bases_future.result()
    total_intents = sum(len(item.intents) for item in amendment_analyses)
    _progress(case_id, f"  amendment docs: {len(amendment_analyses)}")
    _progress(case_id, f"  intents: {total_intents}")
    _progress(case_id, f"  analyzed bases: {len(base_analyses)}")
    _progress(case_id, f"  Время: {_format_seconds(time.perf_counter() - step_analysis_start)}")
    _record_analysis_checks(
        runtime_checklist,
        amendment_analyses=amendment_analyses,
        base_analyses=base_analyses,
    )
    amendment_analyses = _repair_analysis_once(
        runtime_checklist,
        case_id=case_id,
        amendment_analyzer=amendment_analyzer,
        amendment_paths=amendment_paths,
        amendment_analyses=amendment_analyses,
        base_analyses=base_analyses,
    )

    if case_topology == "special_single_amendment_multi_base":
        base_runs = []
        for index, base_doc in enumerate(case.get("base_docs", []), 1):
            _progress(case_id, f"base branch {index}/{len(case.get('base_docs', []))}: {base_doc.name}")
            base_artifacts_dir = artifacts_root / f"base_{index}"
            base_runs.append(
                _run_single_base_flow(
                    case_id=case_id,
                    workspace_root=workspace_root,
                    artifacts_dir=base_artifacts_dir,
                    amendment_analyses=amendment_analyses,
                    base_doc=base_doc,
                    base_analysis=base_analyses[str(base_doc)],
                    runtime_checklist=runtime_checklist,
                    base_analyzer=base_analyzer,
                    skeleton_builder=skeleton_builder,
                    resolver=resolver,
                    editor=editor,
                    checklist_builder=checklist_builder,
                    validator=validator,
                )
            )
        return {
            "case_id": case_id,
            "case_topology": case_topology,
            "workspace_root": str(workspace_root),
            "amendments": [item.to_dict() for item in amendment_analyses],
            "base_runs": base_runs,
            "validation": {
                "is_valid": all(item["validation"]["is_valid"] for item in base_runs),
                "structural_ok": all(item["validation"]["structural_ok"] for item in base_runs),
                "judge_ok": all(item["validation"]["judge_ok"] for item in base_runs),
            },
        }

    base_doc = case.get("base_doc")
    if base_doc is None:
        raise ValueError("Не найден base_doc для single-base topology")
    _progress(case_id, f"single-base flow: {base_doc.name}")
    single_result = _run_single_base_flow(
        case_id=case_id,
        workspace_root=workspace_root,
        artifacts_dir=artifacts_root,
        amendment_analyses=amendment_analyses,
        base_doc=base_doc,
        base_analysis=base_analyses[str(base_doc)],
        runtime_checklist=runtime_checklist,
        base_analyzer=base_analyzer,
        skeleton_builder=skeleton_builder,
        resolver=resolver,
        editor=editor,
        checklist_builder=checklist_builder,
        validator=validator,
    )

    result = {
        "case_id": case_id,
        "case_topology": case_topology,
        "workspace_root": str(workspace_root),
        "amendments": [item.to_dict() for item in amendment_analyses],
        **single_result,
    }
    if len(single_result.get("steps", [])) > 1:
        result["stage_outputs"] = [
            {
                "stage_index": index,
                "output_doc": step["output_doc"],
                "source_document_label": step["amendment"]["metadata"]["document_label"],
            }
            for index, step in enumerate(single_result["steps"], 1)
        ]
    _progress(case_id, "=" * 60)
    _progress(case_id, f"Пайплайн завершён за {_format_seconds(time.perf_counter() - total_start)}")
    _progress(case_id, f"  output: {result.get('output_doc', '')}")
    _progress(
        case_id,
        f"  Валидация: {'OK' if result['validation']['is_valid'] else 'NOT OK'} "
        f"(structural_ok={result['validation']['structural_ok']}, judge_ok={result['validation']['judge_ok']})",
    )
    _progress(case_id, "=" * 60)
    return result


def main() -> None:
    args = parse_args()
    workspace_root = args.workspace_root.resolve()
    case = load_case(workspace_root, args.case_id)
    result = run_case(case, workspace_root, args.models_config.resolve() if args.models_config else None)
    if args.output_json:
        args.output_json.parent.mkdir(parents=True, exist_ok=True)
        args.output_json.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(result["validation"], ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
