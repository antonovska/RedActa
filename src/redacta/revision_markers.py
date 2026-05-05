from __future__ import annotations

from pathlib import Path
import re
from typing import Iterable

from docx import Document

from .base_agent import compact
from .editor_v2 import build_annotation_paragraph
from .schema import ResolvedOperation
from .utils import to_instrumental


class ConsultantMarkerFormatter:
    def format_marker(self, operation: ResolvedOperation) -> str:
        label = to_instrumental(operation.source_document_label)
        if not label:
            return ""

        # Паттерн для repeal_point - "утратил силу. - ..."
        if operation.operation_kind == "repeal_point":
            if operation.point_number:
                return f"{operation.point_number}. Утратил силу. - {label}"
            return f"Утратил силу. - {label}"

        introduced_label = self._introduced_label(operation.source_document_label)

        # Паттерн для replace_table_row
        if operation.operation_kind == "replace_table_row":
            if operation.table_row_ref:
                return f"(строка {operation.table_row_ref} в ред. {label})"
            if operation.structured_entry_ref:
                return f"(запись {operation.structured_entry_ref} в ред. {label})"
            return f"(таблица в ред. {label})"

        # Паттерн для replace_structured_entry
        if operation.operation_kind == "replace_structured_entry":
            if operation.structured_entry_ref:
                return f"(запись {operation.structured_entry_ref} в ред. {label})"
            if operation.table_row_ref:
                return f"(строка {operation.table_row_ref} в ред. {label})"
            return f"(структурированная запись в ред. {label})"

        # Паттерн для replace_person_role
        if operation.operation_kind == "replace_person_role":
            return f"(должность заменена {introduced_label})"

        # Паттерн для insert_list_entry
        if operation.operation_kind == "insert_list_entry":
            return f"(элемент введен {introduced_label})"

        # Паттерн для repeal_appendix_block
        if operation.operation_kind == "repeal_appendix_block":
            if operation.appendix_number:
                return f"(Приложение {operation.appendix_number} утратило силу. - {label})"
            return f"(приложение утратило силу. - {label})"

        # Паттерн для replace_appendix_block
        if operation.operation_kind == "replace_appendix_block":
            if operation.appendix_number:
                return f"(Приложение {operation.appendix_number} в ред. {label})"
            return f"(приложение в ред. {label})"

        # Паттерн для insert_point
        if operation.operation_kind == "insert_point":
            if operation.point_number:
                return f"(п. {operation.point_number} введен {introduced_label})"
            if operation.point_ref:
                return f"(п. {operation.point_ref} введен {introduced_label})"
            return f"(введен {introduced_label})"

        # Паттерн для append_words_to_point
        if operation.operation_kind == "append_words_to_point":
            if operation.point_number:
                return f"(п. {operation.point_number} дополнено {introduced_label})"
            if operation.point_ref:
                return f"(п. {operation.point_ref} дополнено {introduced_label})"
            return f"(дополнено {introduced_label})"

        # Паттерн для абзаца внутри подпункта
        if operation.operation_kind == "append_section_item" and operation.paragraph_ordinal is not None:
            return f"(абзац введен {introduced_label})"

        # Паттерн для замены абзаца внутри подпункта
        if operation.operation_kind == "replace_point" and operation.paragraph_ordinal is not None:
            return f"(в ред. {label})"

        # Паттерн для подпункта (введен/в ред.)
        if operation.subpoint_ref:
            marker = compact(operation.subpoint_ref)
            if operation.operation_kind == "append_section_item":
                return f'(пп. "{marker}" введен {introduced_label})'
            if operation.operation_kind == "replace_point":
                return f'(пп. "{marker}" в ред. {label})'

        # Паттерн для пункта (введен/в ред.)
        if operation.point_ref or operation.point_number:
            if operation.operation_kind == "append_section_item":
                point_num = operation.point_number or operation.point_ref
                return f"(п. {point_num} введен {introduced_label})"
            if operation.operation_kind == "replace_point":
                point_num = operation.point_number or operation.point_ref
                return f"(п. {point_num} в ред. {label})"

        # Общий паттерн для replace_phrase_globally и других операций
        return f"(в ред. {label})"

    def _introduced_label(self, label: str) -> str:
        value = compact(label)

        # Обработка различных типов документов с правильным инструментальным падежом
        prefixes = [
            ("Приказ ", "Приказом "),
            ("Решение ", "Решением "),
            ("Указ ", "Указом "),
            ("Постановление ", "Постановлением "),
            ("Распоряжение ", "Распоряжением "),
            ("Закон ", "Законом "),
            ("Указание ", "Указанием "),
        ]

        for prefix, instrumental in prefixes:
            if value.startswith(prefix):
                return instrumental + value[len(prefix):]

        # Fallback: использовать to_instrumental для необработанных типов
        return to_instrumental(value)


class RevisionMarkerInserter:
    def __init__(self, formatter: ConsultantMarkerFormatter | None = None) -> None:
        self._formatter = formatter or ConsultantMarkerFormatter()

    def insert_markers(self, document_path: Path, operations: Iterable[ResolvedOperation]) -> list[dict[str, object]]:
        document = Document(document_path)
        planned: list[tuple[int, ResolvedOperation, str]] = []
        for operation in operations:
            if operation.status != "resolved" or not operation.paragraph_indices:
                continue
            marker = self._formatter.format_marker(operation)
            if not marker:
                continue
            paragraph_index = self._marker_anchor_index(document, operation)
            if not (0 <= paragraph_index < len(document.paragraphs)):
                continue
            planned.append((paragraph_index, operation, marker))

        inserted: list[dict[str, object]] = []
        seen: set[tuple[int, str]] = set()
        for paragraph_index, operation, marker in sorted(planned, key=lambda item: item[0], reverse=True):
            key = (paragraph_index, marker)
            if key in seen:
                continue
            seen.add(key)
            paragraph = document.paragraphs[paragraph_index]
            if self._next_paragraph_text(paragraph) == marker:
                continue
            paragraph._p.addnext(build_annotation_paragraph(marker))
            inserted.append(
                {
                    "operation_id": operation.operation_id,
                    "paragraph_index": paragraph_index,
                    "marker": marker,
                }
            )
        document.save(document_path)
        return list(reversed(inserted))

    def _marker_anchor_index(self, document: Document, operation: ResolvedOperation) -> int:
        paragraph_index = operation.paragraph_indices[0]
        if operation.operation_kind != "replace_phrase_globally":
            return paragraph_index
        return self._structural_block_end_index(document, paragraph_index)

    def _structural_block_end_index(self, document: Document, paragraph_index: int) -> int:
        if not (0 <= paragraph_index < len(document.paragraphs)):
            return paragraph_index
        top_point_pattern = re.compile(r"^\d+\.\s+")
        subpoint_pattern = re.compile(r"^[а-яёa-z](?:\(\d+\))?\)\s+", re.IGNORECASE)
        paragraphs = document.paragraphs
        start_index = paragraph_index
        text = paragraphs[start_index].text.strip()
        if not top_point_pattern.match(text) and not subpoint_pattern.match(text):
            for idx in range(start_index - 1, -1, -1):
                previous_text = paragraphs[idx].text.strip()
                if not previous_text:
                    continue
                if top_point_pattern.match(previous_text) or subpoint_pattern.match(previous_text):
                    start_index = idx
                break

        end_index = paragraph_index
        for idx in range(paragraph_index + 1, len(paragraphs)):
            text = paragraphs[idx].text.strip()
            if not text:
                continue
            if top_point_pattern.match(text) or subpoint_pattern.match(text):
                break
            end_index = idx
        return end_index

    def _next_paragraph_text(self, paragraph) -> str:
        next_element = paragraph._p.getnext()
        if next_element is None:
            return ""
        text_parts = [
            node.text
            for node in next_element.iter()
            if node.tag.endswith("}t") and node.text
        ]
        return compact("".join(text_parts))
