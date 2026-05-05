from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document

from .base_agent import compact, is_heading_continuation, is_section_heading, tokenize


@dataclass
class ParagraphRecord:
    absolute_index: int
    text: str


@dataclass
class HeadingGroup:
    text: str
    start_index: int
    end_index: int


def read_paragraph_records(doc_path: Path) -> list[ParagraphRecord]:
    document = Document(doc_path)
    records: list[ParagraphRecord] = []
    for idx, paragraph in enumerate(document.paragraphs):
        text = compact(paragraph.text)
        if text:
            records.append(ParagraphRecord(absolute_index=idx, text=text))
    return records


def build_heading_groups(records: list[ParagraphRecord]) -> list[HeadingGroup]:
    groups: list[HeadingGroup] = []
    idx = 0
    while idx < len(records):
        current = records[idx]
        if not is_section_heading(current.text):
            idx += 1
            continue
        parts = [current.text]
        start = current.absolute_index
        end = current.absolute_index
        j = idx + 1
        while j < len(records) and is_heading_continuation(parts[-1], records[j].text):
            parts.append(records[j].text)
            end = records[j].absolute_index
            j += 1
        groups.append(HeadingGroup(text=" ".join(parts), start_index=start, end_index=end))
        idx = j
    return groups


def find_point_paragraph(records: list[ParagraphRecord], point_number: int) -> ParagraphRecord | None:
    pattern = re.compile(rf"^{point_number}\.\s+")
    for record in records:
        if pattern.match(record.text):
            return record
    return None


def find_all_point_paragraphs(records: list[ParagraphRecord], point_number: int) -> list[ParagraphRecord]:
    """Возвращает ВСЕ абзацы-заголовки пункта N из всего документа.

    Нужно когда в документе несколько нумерованных секций с одинаковыми номерами
    (например, «3.» в теле приказа и «3.» в тексте Порядка-приложения).
    """
    pattern = re.compile(rf"^{point_number}\.\s+")
    return [record for record in records if pattern.match(record.text)]


_SUBPOINT_LETTER_RE = re.compile(r"^([а-яёa-z](?:\(\d+\))?)\)\s+", re.IGNORECASE)


def find_point_block(records: list[ParagraphRecord], point_number: int) -> list[ParagraphRecord]:
    """Возвращает все абзацы верхнеуровневого пункта N (включая заголовок).

    Пункт начинается с «N. текст» и заканчивается перед следующим таким же пунктом
    или концом записей.
    """
    top_point_re = re.compile(r"^(\d+)\.\s+")
    start_idx: int | None = None
    for i, rec in enumerate(records):
        m = top_point_re.match(rec.text)
        if m and int(m.group(1)) == point_number:
            start_idx = i
            break
    if start_idx is None:
        return []
    end_idx = len(records)
    for i in range(start_idx + 1, len(records)):
        if top_point_re.match(records[i].text):
            end_idx = i
            break
    return records[start_idx:end_idx]


def find_all_point_blocks(records: list[ParagraphRecord], point_number: int) -> list[list[ParagraphRecord]]:
    """Возвращает ВСЕ блоки пункта N из документа.

    Нужно когда в документе несколько нумерованных секций с одинаковыми номерами
    (например, «2.» в теле приказа и «2.» в тексте Порядка-приложения).
    """
    top_point_re = re.compile(r"^(\d+)\.\s+")
    blocks: list[list[ParagraphRecord]] = []
    start_idx: int | None = None
    for i, rec in enumerate(records):
        m = top_point_re.match(rec.text)
        if m and int(m.group(1)) == point_number:
            if start_idx is not None:
                # закрываем предыдущий найденный блок
                end_idx = i
                blocks.append(records[start_idx:end_idx])
            start_idx = i
        elif m and start_idx is not None:
            # следующий пункт с другим номером — закрываем текущий
            end_idx = i
            blocks.append(records[start_idx:end_idx])
            start_idx = None
    if start_idx is not None:
        blocks.append(records[start_idx:])
    return blocks


def find_subpoint_in_point(
    point_records: list[ParagraphRecord],
    subpoint_ref: str,
) -> list[ParagraphRecord]:
    """Внутри блока пункта находит все абзацы подпункта subpoint_ref.

    Подпункты форматированы как «б) текст». Блок подпункта заканчивается перед
    следующим подпунктом того же уровня или концом блока пункта.
    Если подпункт не найден — возвращает [].
    """
    target_re = re.compile(rf"^{re.escape(subpoint_ref)}\)\s+", re.IGNORECASE)
    sub_start: int | None = None
    for i, rec in enumerate(point_records):
        if target_re.match(rec.text):
            sub_start = i
            break
    if sub_start is None:
        return []
    sub_end = len(point_records)
    for i in range(sub_start + 1, len(point_records)):
        if _SUBPOINT_LETTER_RE.match(point_records[i].text):
            sub_end = i
            break
    return point_records[sub_start:sub_end]


def find_last_subpoint_in_point(point_records: list[ParagraphRecord]) -> ParagraphRecord | None:
    """Возвращает последний абзац последнего подпункта в блоке пункта.

    Полезно для операций вставки нового подпункта в конец пункта.
    """
    last_sub_idx: int | None = None
    for i, rec in enumerate(point_records):
        if _SUBPOINT_LETTER_RE.match(rec.text):
            last_sub_idx = i
    if last_sub_idx is None:
        return None
    sub_end = len(point_records)
    for i in range(last_sub_idx + 1, len(point_records)):
        if _SUBPOINT_LETTER_RE.match(point_records[i].text):
            sub_end = i
            break
    return point_records[sub_end - 1]


def find_point_ref_paragraph(records: list[ParagraphRecord], point_ref: str) -> ParagraphRecord | None:
    if not point_ref:
        return None
    pattern = re.compile(rf"^{re.escape(point_ref)}\.\s+")
    for record in records:
        if pattern.match(record.text):
            return record
    return None


def find_section_candidates(records: list[ParagraphRecord], section_hint: str) -> list[dict[str, Any]]:
    if not section_hint:
        return []
    headings = build_heading_groups(records)
    hint_tokens = set(tokenize(section_hint))
    if not hint_tokens:
        return []

    candidates: list[dict[str, Any]] = []
    for idx, heading in enumerate(headings):
        heading_tokens = set(tokenize(heading.text))
        overlap = len(hint_tokens.intersection(heading_tokens))
        if overlap == 0:
            continue
        next_start = headings[idx + 1].start_index if idx + 1 < len(headings) else 10 ** 9
        section_records = [
            record for record in records
            if record.absolute_index > heading.end_index and record.absolute_index < next_start
        ]
        item_records = [record for record in section_records if not is_section_heading(record.text)]
        if not item_records:
            continue
        candidates.append(
            {
                "heading_text": heading.text,
                "heading_start_index": heading.start_index,
                "overlap": overlap,
                "item_records": item_records,
            }
        )
    candidates.sort(key=lambda item: item["overlap"], reverse=True)
    return candidates


def find_appendix_start(records: list[ParagraphRecord], appendix_number: str) -> ParagraphRecord | None:
    pattern = re.compile(rf"^приложение\s+n?\s*{re.escape(appendix_number)}(?:\b|$)", re.IGNORECASE)
    for record in records:
        if pattern.match(record.text):
            return record
    return None


def find_table_section_candidates(doc_path: Path, section_hint: str) -> list[dict[str, Any]]:
    if not section_hint:
        return []
    hint_tokens = set(tokenize(section_hint))
    if not hint_tokens:
        return []

    document = Document(doc_path)
    candidates: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables):
        section_rows: list[tuple[int, str]] = []
        for row_index, row in enumerate(table.rows):
            values = [compact(cell.text) for cell in row.cells]
            non_empty = [value for value in values if value]
            if not non_empty:
                continue
            unique = list(dict.fromkeys(non_empty))
            if len(unique) == 1:
                section_rows.append((row_index, unique[0]))

        if not section_rows:
            continue

        for idx, (section_row_index, section_text) in enumerate(section_rows):
            section_tokens = set(tokenize(section_text))
            overlap = len(hint_tokens.intersection(section_tokens))
            if overlap == 0:
                continue
            next_section_row_index = section_rows[idx + 1][0] if idx + 1 < len(section_rows) else len(table.rows)
            data_rows = [row_idx for row_idx in range(section_row_index + 1, next_section_row_index)]
            if not data_rows:
                continue
            candidates.append(
                {
                    "table_index": table_index,
                    "heading_text": section_text,
                    "anchor_row_index": data_rows[-1],
                    "overlap": overlap,
                }
            )
    candidates.sort(key=lambda item: item["overlap"], reverse=True)
    return candidates
