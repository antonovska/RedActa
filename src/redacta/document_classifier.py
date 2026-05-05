"""Document complexity classifier.

Classifies amendment and base documents into one of three complexity tiers
before the pipeline starts processing, enabling early routing decisions.

Tiers
-----
plain
    Document contains only paragraphs and simple inline text.
    Standard pipeline path applies.

table_heavy
    Document contains tables with structured data rows (row\\t... payload).
    Table-aware resolver and editor paths apply.

media_heavy
    Document contains embedded images, drawings, or formulas (w:drawing,
    w:object, etc.).  These require manual review or a dedicated media path
    (L3-03 in the backlog).
"""
from __future__ import annotations

from pathlib import Path
from typing import Literal

from docx import Document
from docx.oxml.ns import qn

DocumentComplexity = Literal["plain", "table_heavy", "media_heavy"]


def classify_amendment_complexity(doc_path: Path) -> DocumentComplexity:
    """Classify an amendment document by its structural complexity.

    Priority: media_heavy > table_heavy > plain.
    """
    try:
        document = Document(doc_path)
    except Exception:
        return "plain"

    # Check for embedded media in paragraphs
    if _has_drawings(document):
        return "media_heavy"

    # Check for tables with actual data rows (not just header/section rows)
    if _has_data_table(document):
        return "table_heavy"

    return "plain"


def classify_base_complexity(doc_path: Path) -> DocumentComplexity:
    """Classify a base document by its structural complexity.

    Used to anticipate what kind of edits will be needed.
    Priority: media_heavy > table_heavy > plain.
    """
    try:
        document = Document(doc_path)
    except Exception:
        return "plain"

    if _has_drawings(document):
        return "media_heavy"

    if _has_data_table(document):
        return "table_heavy"

    return "plain"


def _has_drawings(document: Document) -> bool:
    """Return True if any paragraph contains an embedded drawing or object."""
    drawing_tag = qn("w:drawing")
    object_tag = qn("w:object")
    for paragraph in document.paragraphs:
        for child in paragraph._p:
            if child.tag in (drawing_tag, object_tag):
                return True
    # Also check inside table cells
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for child in paragraph._p:
                        if child.tag in (drawing_tag, object_tag):
                            return True
    return False


def _has_data_table(document: Document) -> bool:
    """Return True if the document contains a table with at least one data row.

    A data row is a row where cells have distinct non-empty values
    (as opposed to a merged section-header row where all cells repeat the same text).
    """
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            non_empty = [c for c in cells if c]
            if not non_empty:
                continue
            unique = list(dict.fromkeys(non_empty))
            # More than one distinct value → data row
            if len(unique) > 1:
                return True
    return False
