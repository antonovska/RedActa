from __future__ import annotations

from docx import Document

from redacta.editor_v2 import EditorV2, normalize_item_text
from redacta.schema import ResolvedOperation


def test_normalize_item_text_strips_outer_revision_quotes() -> None:
    assert normalize_item_text('"1.1. New text."') == "1.1. New text."
    assert normalize_item_text('"1.1. New text".') == "1.1. New text."


def test_replace_point_prefers_decimal_point_ref_over_integer_point_number(tmp_path) -> None:
    input_doc = tmp_path / "input.docx"
    output_doc = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("1. Section heading")
    document.add_paragraph("1.1. Old first subpoint.")
    document.add_paragraph("1.2. Old second subpoint.")
    document.save(input_doc)
    operation = ResolvedOperation(
        operation_id="c1",
        operation_kind="replace_point",
        status="resolved",
        source_document_label="Распоряжение ОАО РЖД от 28.04.2020 N 944 р",
        paragraph_indices=[1],
        point_ref="1.1",
        point_number=1,
        new_text='"1.1. New first subpoint."',
    )

    EditorV2().edit(input_doc, output_doc, [operation])
    paragraphs = [" ".join(paragraph.text.split()) for paragraph in Document(output_doc).paragraphs if paragraph.text.strip()]

    assert paragraphs == [
        "1. Section heading",
        "1.1. New first subpoint.",
        "1.2. Old second subpoint.",
    ]
