from __future__ import annotations

from graph_pipeline.revision_markers import ConsultantMarkerFormatter, RevisionMarkerInserter
from graph_pipeline.editor_v2 import EditorV2, normalize_item_text
from graph_pipeline.schema import ResolvedOperation
from docx import Document


def _operation(kind: str, **kwargs: object) -> ResolvedOperation:
    return ResolvedOperation(
        operation_id="c1",
        operation_kind=kind,
        status="resolved",
        source_document_label="Приказ Минфина России от 03.10.2025 N 141н",
        **kwargs,
    )


def test_formats_subpoint_replacement_marker() -> None:
    marker = ConsultantMarkerFormatter().format_marker(
        _operation("replace_point", parent_point_ref="2", subpoint_ref="б")
    )

    assert marker == '(пп. "б" в ред. Приказа Минфина России от 03.10.2025 N 141н)'


def test_formats_inserted_subpoint_marker() -> None:
    marker = ConsultantMarkerFormatter().format_marker(
        _operation("append_section_item", parent_point_ref="2", subpoint_ref="е(1)")
    )

    assert marker == '(пп. "е(1)" введен Приказом Минфина России от 03.10.2025 N 141н)'


def test_formats_inserted_paragraph_marker() -> None:
    marker = ConsultantMarkerFormatter().format_marker(
        _operation("append_section_item", parent_point_ref="4", subpoint_ref="в", paragraph_ordinal=8)
    )

    assert marker == "(абзац введен Приказом Минфина России от 03.10.2025 N 141н)"


def test_formatter_ignores_legacy_note_text_for_consultant_marker() -> None:
    marker = ConsultantMarkerFormatter().format_marker(
        _operation(
            "replace_point",
            parent_point_ref="2",
            subpoint_ref="б",
            note_text="(подп. б п. 2 в ред. legacy)",
        )
    )

    assert marker == '(пп. "б" в ред. Приказа Минфина России от 03.10.2025 N 141н)'


def test_repeal_point_has_no_separate_revision_marker() -> None:
    marker = ConsultantMarkerFormatter().format_marker(
        _operation("repeal_point", point_ref="3")
    )

    assert marker == ""


def test_paragraph_replacement_uses_generic_revision_marker() -> None:
    marker = ConsultantMarkerFormatter().format_marker(
        _operation("replace_point", parent_point_ref="4", subpoint_ref="в", paragraph_ordinal=7)
    )

    assert marker == "(в ред. Приказа Минфина России от 03.10.2025 N 141н)"


def test_structural_editor_does_not_insert_revision_markers(tmp_path) -> None:
    input_doc = tmp_path / "input.docx"
    output_doc = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("б) old;")
    document.save(input_doc)
    operation = _operation(
        "replace_point",
        paragraph_indices=[0],
        subpoint_ref="б",
        new_text="б) new;",
        note_text='(пп. "б" в ред. Приказа Минфина России от 03.10.2025 N 141н)',
    )

    result = EditorV2().edit(input_doc, output_doc, [operation])
    paragraphs = [" ".join(paragraph.text.split()) for paragraph in Document(output_doc).paragraphs if paragraph.text.strip()]

    assert paragraphs == ["б) new;"]
    assert result["applied_operations"][0]["operation_id"] == "c1"
    assert result["applied_operations"][0]["paragraph_indices"] == [0]


def test_revision_marker_inserter_adds_marker_after_structural_edit(tmp_path) -> None:
    docx_path = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("б) new;")
    document.save(docx_path)
    operation = _operation(
        "replace_point",
        paragraph_indices=[0],
        subpoint_ref="б",
    )

    inserted = RevisionMarkerInserter().insert_markers(docx_path, [operation])
    paragraphs = [" ".join(paragraph.text.split()) for paragraph in Document(docx_path).paragraphs if paragraph.text.strip()]

    assert paragraphs == [
        "б) new;",
        '(пп. "б" в ред. Приказа Минфина России от 03.10.2025 N 141н)',
    ]
    assert inserted == [
        {
            "operation_id": "c1",
            "paragraph_index": 0,
            "marker": '(пп. "б" в ред. Приказа Минфина России от 03.10.2025 N 141н)',
        }
    ]


def test_editor_updates_final_anchor_for_marker_post_pass(tmp_path) -> None:
    input_doc = tmp_path / "input.docx"
    output_doc = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("а) previous;")
    document.add_paragraph("б) target;")
    document.add_paragraph("в) later;")
    document.save(input_doc)
    operation = _operation(
        "append_section_item",
        insert_after_index=1,
        parent_point_ref="2",
        subpoint_ref="б(1)",
        new_item_text="б(1)) inserted;",
    )

    result = EditorV2().edit(input_doc, output_doc, [operation])

    assert result["applied_operations"][0]["paragraph_indices"] == [2]


def test_editor_rebases_previous_operation_anchor_after_later_insert(tmp_path) -> None:
    input_doc = tmp_path / "input.docx"
    output_doc = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("anchor;")
    document.add_paragraph("middle;")
    document.add_paragraph("old target;")
    document.save(input_doc)
    phrase_operation = ResolvedOperation(
        operation_id="c1",
        operation_kind="replace_phrase_globally",
        status="resolved",
        source_document_label="\u041f\u0440\u0438\u043a\u0430\u0437 \u041c\u0438\u043d\u0444\u0438\u043d\u0430 \u0420\u043e\u0441\u0441\u0438\u0438 \u043e\u0442 03.10.2025 N 141\u043d",
        paragraph_indices=[2],
        old_text="old",
        new_text="new",
    )
    insert_operation = ResolvedOperation(
        operation_id="c2",
        operation_kind="append_section_item",
        status="resolved",
        source_document_label="\u041f\u0440\u0438\u043a\u0430\u0437 \u041c\u0438\u043d\u0444\u0438\u043d\u0430 \u0420\u043e\u0441\u0441\u0438\u0438 \u043e\u0442 03.10.2025 N 141\u043d",
        insert_after_index=0,
        new_item_text="inserted;",
    )

    result = EditorV2().edit(input_doc, output_doc, [phrase_operation, insert_operation])
    applied = {operation["operation_id"]: operation for operation in result["applied_operations"]}

    assert applied["c1"]["paragraph_indices"] == [3]
    assert applied["c2"]["paragraph_indices"] == [1]


def test_phrase_replacement_marker_uses_structural_block_end(tmp_path) -> None:
    docx_path = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("2. Items:")
    document.add_paragraph("\u0430) changed heading.")
    document.add_paragraph("changed continuation.")
    document.add_paragraph("\u0431) next.")
    document.save(docx_path)
    operations = [
        _operation("replace_phrase_globally", paragraph_indices=[1]),
        _operation("replace_phrase_globally", paragraph_indices=[2]),
    ]

    inserted = RevisionMarkerInserter().insert_markers(docx_path, operations)
    paragraphs = [" ".join(paragraph.text.split()) for paragraph in Document(docx_path).paragraphs if paragraph.text.strip()]

    marker = "(\u0432 \u0440\u0435\u0434. \u041f\u0440\u0438\u043a\u0430\u0437\u0430 \u041c\u0438\u043d\u0444\u0438\u043d\u0430 \u0420\u043e\u0441\u0441\u0438\u0438 \u043e\u0442 03.10.2025 N 141\u043d)"
    assert paragraphs == [
        "2. Items:",
        "\u0430) changed heading.",
        "changed continuation.",
        marker,
        "\u0431) next.",
    ]
    assert inserted == [
        {
            "operation_id": "c1",
            "paragraph_index": 2,
            "marker": marker,
        }
    ]


def test_editor_appends_paragraph_inside_existing_subpoint(tmp_path) -> None:
    input_doc = tmp_path / "input.docx"
    output_doc = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("4. Documents:")
    document.add_paragraph("\u0430) first;")
    document.add_paragraph("\u0431) second;")
    document.add_paragraph("\u0432) target:")
    document.add_paragraph("existing continuation;")
    document.add_paragraph("\u0433) next.")
    document.save(input_doc)
    operation = _operation(
        "append_section_item",
        parent_point_number=4,
        parent_point_ref="4",
        subpoint_ref="\u0432",
        new_item_text="new continuation.",
    )

    result = EditorV2().edit(input_doc, output_doc, [operation])
    paragraphs = [" ".join(paragraph.text.split()) for paragraph in Document(output_doc).paragraphs if paragraph.text.strip()]

    assert paragraphs == [
        "4. Documents:",
        "\u0430) first;",
        "\u0431) second;",
        "\u0432) target:",
        "existing continuation;",
        "new continuation.",
        "\u0433) next.",
    ]
    assert result["applied_operations"][0]["paragraph_indices"] == [5]
    assert result["applied_operations"][0]["paragraph_ordinal"] == 1


def test_normalize_item_text_removes_duplicate_terminal_punctuation() -> None:
    assert normalize_item_text("text;;") == "text;"
    assert normalize_item_text('text.".') == "text."
