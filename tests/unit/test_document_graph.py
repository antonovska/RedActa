import sys
import unittest
from io import BytesIO
from pathlib import Path

from docx import Document


sys.path.insert(0, str(Path(__file__).resolve().parents[2] / "src"))

from graph_pipeline.document_graph import DocumentGraph


def _make_docx() -> BytesIO:
    doc = Document()
    doc.add_paragraph("1. General provisions")
    doc.add_paragraph("1.1. Subpoint")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "N"
    table.cell(0, 1).text = "Name"
    table.cell(1, 0).text = "4"
    table.cell(1, 1).text = "Old row"
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


class DocumentGraphTest(unittest.TestCase):
    def test_builds_paragraph_and_table_nodes(self) -> None:
        graph = DocumentGraph.from_docx(_make_docx())
        self.assertTrue(graph.find_by_type("paragraph"))
        self.assertTrue(graph.find_by_type("table"))
        self.assertTrue(graph.find_by_type("row"))
        self.assertTrue(graph.find_by_type("cell"))

    def test_finds_table_row_by_first_cell_text(self) -> None:
        graph = DocumentGraph.from_docx(_make_docx())
        rows = graph.find_table_rows_by_ref("4")
        self.assertEqual(1, len(rows))
        self.assertEqual("row", rows[0].node_type)
        self.assertIn("Old row", rows[0].text)


if __name__ == "__main__":
    unittest.main()
