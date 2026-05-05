import sys
import types
import uuid
from pathlib import Path
import unittest

from docx import Document


sys.path.insert(0, str(Path(__file__).resolve().parents[2] / "src"))
sys.modules.setdefault("openai", types.SimpleNamespace(OpenAI=object))

from graph_pipeline.editor_v2 import EditorV2
from graph_pipeline.schema import ResolvedOperation


class EditorV2DriftTest(unittest.TestCase):
    def test_phrase_replacement_does_not_insert_marker_or_shift_following_repeal(self) -> None:
        tmp_path = Path(__file__).resolve().parents[2] / "tmp" / "tests" / uuid.uuid4().hex
        tmp_path.mkdir(parents=True, exist_ok=True)
        input_doc = tmp_path / "input.docx"
        output_doc = tmp_path / "output.docx"

        document = Document()
        document.add_paragraph("preamble old citation text")
        document.add_paragraph("1. keep this point")
        document.add_paragraph("2. target point")
        document.add_paragraph("3. later point")
        document.save(input_doc)

        operations = [
            ResolvedOperation(
                operation_id="c1",
                operation_kind="replace_phrase_globally",
                status="resolved",
                source_document_label="source",
                paragraph_indices=[0],
                old_text="old citation",
                new_text="",
                note_text="(edited)",
            ),
            ResolvedOperation(
                operation_id="c2",
                operation_kind="repeal_point",
                status="resolved",
                source_document_label="source",
                paragraph_indices=[2],
                point_ref="2",
                point_number=2,
                new_text="2. repealed",
            ),
        ]

        result = EditorV2().edit(input_doc, output_doc, operations)

        texts = [paragraph.text for paragraph in Document(output_doc).paragraphs]
        self.assertEqual([], result["drift_events"])
        self.assertEqual("preamble  text", texts[0])
        self.assertEqual("1. keep this point", texts[1])
        self.assertEqual("2. repealed", texts[2])
        self.assertEqual("3. later point", texts[3])

    def test_append_words_updates_semicolon_item_without_marker_drift(self) -> None:
        tmp_path = Path(__file__).resolve().parents[2] / "tmp" / "tests" / uuid.uuid4().hex
        tmp_path.mkdir(parents=True, exist_ok=True)
        input_doc = tmp_path / "input.docx"
        output_doc = tmp_path / "output.docx"

        document = Document()
        document.add_paragraph("а) previous;")
        document.add_paragraph("б) target words;")
        document.add_paragraph("в) later;")
        document.save(input_doc)

        operations = [
            ResolvedOperation(
                operation_id="c1",
                operation_kind="append_words_to_point",
                status="resolved",
                source_document_label="source",
                paragraph_indices=[1],
                point_ref="б",
                appended_words=", appended words",
                note_text="(edited)",
            ),
            ResolvedOperation(
                operation_id="c2",
                operation_kind="repeal_point",
                status="resolved",
                source_document_label="source",
                paragraph_indices=[2],
                point_ref="в",
                new_text="в) repealed;",
            ),
        ]

        result = EditorV2().edit(input_doc, output_doc, operations)

        texts = [paragraph.text for paragraph in Document(output_doc).paragraphs]
        self.assertEqual([], result["drift_events"])
        self.assertEqual("б) target words, appended words;", texts[1])
        self.assertEqual("в) repealed;", texts[2])

    def test_repeal_point_does_not_add_duplicate_revision_marker(self) -> None:
        tmp_path = Path(__file__).resolve().parents[2] / "tmp" / "tests" / uuid.uuid4().hex
        tmp_path.mkdir(parents=True, exist_ok=True)
        input_doc = tmp_path / "input.docx"
        output_doc = tmp_path / "output.docx"

        document = Document()
        document.add_paragraph("1. keep")
        document.add_paragraph("2. target")
        document.add_paragraph("3. later")
        document.save(input_doc)

        operations = [
            ResolvedOperation(
                operation_id="c1",
                operation_kind="repeal_point",
                status="resolved",
                source_document_label="source",
                paragraph_indices=[1],
                point_ref="2",
                point_number=2,
                new_text="2. Repealed. - source.",
            )
        ]

        EditorV2().edit(input_doc, output_doc, operations)

        texts = [paragraph.text for paragraph in Document(output_doc).paragraphs]
        self.assertEqual(["1. keep", "2. Repealed. - source.", "3. later"], texts)

    def test_append_new_subpoint_uses_current_previous_subpoint_anchor(self) -> None:
        tmp_path = Path(__file__).resolve().parents[2] / "tmp" / "tests" / uuid.uuid4().hex
        tmp_path.mkdir(parents=True, exist_ok=True)
        input_doc = tmp_path / "input.docx"
        output_doc = tmp_path / "output.docx"

        document = Document()
        document.add_paragraph("2. parent point")
        document.add_paragraph("д) replaced d;")
        document.add_paragraph("(note d)")
        document.add_paragraph("е) replaced e;")
        document.add_paragraph("(note e)")
        document.add_paragraph("ж) next subpoint.")
        document.save(input_doc)

        operations = [
            ResolvedOperation(
                operation_id="c1",
                operation_kind="append_section_item",
                status="resolved",
                source_document_label="source",
                insert_after_index=1,
                point_ref="е(1)",
                parent_point_number=2,
                subpoint_ref="е(1)",
                new_item_text="е(1)) inserted subpoint;",
                note_text="(inserted)",
            )
        ]

        EditorV2().edit(input_doc, output_doc, operations)

        texts = [paragraph.text for paragraph in Document(output_doc).paragraphs]
        self.assertEqual(
            [
                "2. parent point",
                "д) replaced d;",
                "(note d)",
                "е) replaced e;",
                "(note e)",
                "е(1)) inserted subpoint;",
                "ж) next subpoint.",
            ],
            texts,
        )


if __name__ == "__main__":
    unittest.main()
