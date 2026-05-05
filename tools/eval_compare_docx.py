from __future__ import annotations

import argparse
import json
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any

from docx import Document


def _compact(text: str) -> str:
    return " ".join((text or "").replace("\xa0", " ").split())


def extract_docx_text(path: Path) -> dict[str, Any]:
    document = Document(path)
    paragraphs = [_compact(paragraph.text) for paragraph in document.paragraphs if _compact(paragraph.text)]
    tables: list[list[list[str]]] = []
    for table in document.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([_compact(cell.text) for cell in row.cells])
        tables.append(rows)
    return {"paragraphs": paragraphs, "tables": tables}


def compare_docx_text(result_path: Path, reference_path: Path) -> dict[str, Any]:
    result = extract_docx_text(result_path)
    reference = extract_docx_text(reference_path)
    matcher = SequenceMatcher(
        a=reference["paragraphs"],
        b=result["paragraphs"],
        autojunk=False,
    )
    changes = [item for item in matcher.get_opcodes() if item[0] != "equal"]
    return {
        "result_path": str(result_path),
        "reference_path": str(reference_path),
        "result_paragraphs": len(result["paragraphs"]),
        "reference_paragraphs": len(reference["paragraphs"]),
        "result_tables": len(result["tables"]),
        "reference_tables": len(reference["tables"]),
        "diff_blocks": len(changes),
        "table_equal": result["tables"] == reference["tables"],
        "changes": [
            {
                "tag": tag,
                "reference_range": [i1, i2],
                "result_range": [j1, j2],
                "reference_preview": reference["paragraphs"][i1:i2][:3],
                "result_preview": result["paragraphs"][j1:j2][:3],
            }
            for tag, i1, i2, j1, j2 in changes
        ],
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Dev-only DOCX text comparison helper")
    parser.add_argument("result_docx", type=Path)
    parser.add_argument("reference_docx", type=Path)
    args = parser.parse_args()
    print(json.dumps(compare_docx_text(args.result_docx, args.reference_docx), ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
