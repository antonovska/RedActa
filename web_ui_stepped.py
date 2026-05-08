# RedActa — Interactive Step-by-Step Pipeline Explorer
# Streamlit-based UI for running and inspecting each pipeline stage individually.
#
# Usage:
#   cd C:\Users\appan\Downloads\RedActa_1\RedActa-main
#   pip install streamlit
#   streamlit run web_ui_stepped.py

from __future__ import annotations

import json
import shutil
import tempfile
import time
import traceback
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from typing import Any

import streamlit as st

# Ensure src is on path
import sys

sys.path.insert(0, str(Path(__file__).parent / "src"))

from docx import Document

from redacta.amendment_analyzer import AmendmentAnalyzer
from redacta.base_analyzer import BaseAnalyzer
from redacta.case_loader import load_case
from redacta.config import load_models_config
from redacta.editor_v2 import EditorV2
from redacta.manual_review import split_operations_for_manual_review
from redacta.pipeline_checklist import PipelineChecklist
from redacta.resolver import PipelineResolver
from redacta.revision_markers import RevisionMarkerInserter
from redacta.run_case import (
    _analysis_blocked_operations,
    _analysis_gate_details,
    _enforce_resolution_gate,
    _extract_amendment_directives,
    _manual_review_result,
    _manual_review_validation,
    _record_analysis_checks,
    _record_initial_checks,
    _repair_analysis_once,
)
from redacta.service_tables import build_service_table_specs
from redacta.skeleton_builder import SkeletonBuilder
from redacta.validation_checklist_builder import ValidationChecklistBuilder
from redacta.validator import StrictJudgeValidator

st.set_page_config(
    page_title="RedActa Step-by-Step",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ---------------------------------------------------------------------------
# CSS
# ---------------------------------------------------------------------------
st.markdown(
    """
    <style>
    .pipeline-map {
        display: flex;
        flex-direction: column;
        gap: 0.5rem;
        margin-bottom: 1.5rem;
    }
    .pipeline-row {
        display: flex;
        gap: 0.5rem;
        justify-content: center;
        align-items: stretch;
    }
    .step-node {
        flex: 1;
        max-width: 220px;
        min-height: 60px;
        border-radius: 8px;
        padding: 0.6rem 0.8rem;
        font-size: 0.85rem;
        font-weight: 600;
        display: flex;
        flex-direction: column;
        justify-content: center;
        align-items: center;
        text-align: center;
        transition: all 0.2s ease;
        border: 2px solid transparent;
        cursor: default;
    }
    .step-node .step-label {
        font-size: 0.75rem;
        font-weight: 400;
        opacity: 0.85;
        margin-top: 2px;
    }
    .step-pending {
        background-color: #f3f4f6;
        color: #9ca3af;
        border-color: #e5e7eb;
    }
    .step-current {
        background-color: #dbeafe;
        color: #1e40af;
        border-color: #3b82f6;
        box-shadow: 0 0 0 3px rgba(59,130,246,0.15);
    }
    .step-success {
        background-color: #dcfce7;
        color: #166534;
        border-color: #22c55e;
    }
    .step-warning {
        background-color: #fef3c7;
        color: #92400e;
        border-color: #f59e0b;
    }
    .step-error {
        background-color: #fee2e2;
        color: #991b1b;
        border-color: #ef4444;
    }
    .arrow-down {
        text-align: center;
        color: #9ca3af;
        font-size: 1.2rem;
        line-height: 0.8;
    }
    .result-card {
        background: #ffffff;
        border: 1px solid #e5e7eb;
        border-radius: 10px;
        padding: 1.25rem;
        margin-top: 1rem;
    }
    .log-line {
        font-family: "SFMono-Regular", Consolas, "Liberation Mono", Menlo, monospace;
        font-size: 0.8rem;
        line-height: 1.4;
        color: #374151;
        white-space: pre-wrap;
        word-break: break-word;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def get_default_models_config() -> dict[str, Any]:
    return {
        "runtime": {
            "base_url": "http://127.0.0.1:11434/v1",
            "api_key": "ollama",
            "temperature": 0.0,
            "timeout": 900,
            "max_retries": 4,
            "enable_thinking": False,
            "semantic_ranking_enabled": False,
            "embedding_provider": "local_http",
            "embedding_service_url": "http://127.0.0.1:8010",
            "embedding_model_path": "models/deepvk_USER2-base",
            "semantic_top_k": 5,
            "semantic_auto_threshold": 0.72,
            "semantic_auto_margin": 0.08,
        },
        "models": {
            "default": "hf.co/ai-sage/GigaChat3.1-10B-A1.8B-GGUF:Q4_K_M",
            "analyst": "hf.co/ai-sage/GigaChat3.1-10B-A1.8B-GGUF:Q4_K_M",
            "resolver_disambiguation": "hf.co/ai-sage/GigaChat3.1-10B-A1.8B-GGUF:Q4_K_M",
            "validator": "hf.co/ai-sage/GigaChat3.1-10B-A1.8B-GGUF:Q4_K_M",
        },
    }


@st.cache_data(show_spinner=False)
def extract_docx_preview(file_path: str | Path, max_paras: int = 50) -> str:
    try:
        document = Document(str(file_path))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        preview = "\n".join(paragraphs[:max_paras])
        if len(paragraphs) > max_paras:
            preview += f"\n\n... ({len(paragraphs) - max_paras} more paragraphs)"
        return preview
    except Exception as exc:
        return f"Error reading document: {exc}"


def _ensure_state() -> None:
    defaults = {
        "files_ready": False,
        "step_index": 0,
        "step_status": {},  # step_key -> "pending" | "running" | "success" | "warning" | "error"
        "step_results": {},
        "step_logs": {},
        "case": None,
        "workspace": None,
        "models_config_path": None,
        "agents": {},
        "checklist": None,
        "base_doc": None,
        "amendment_paths": None,
        "amendment_analyses": None,
        "base_analysis": None,
        "skeleton_doc": None,
        "service_table_specs": None,
        "current_doc": None,
        "resolution": None,
        "split": None,
        "edit_result": None,
        "marker_result": None,
        "final_base_analysis": None,
        "final_specs": None,
        "validation_checklist": None,
        "validation": None,
        "blocked_operations": [],
        "all_operations": [],
        "all_statuses": [],
        "pipeline_error": None,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


def _set_step_status(step: str, status: str) -> None:
    st.session_state.step_status[step] = status


def _add_log(step: str, message: str) -> None:
    st.session_state.step_logs.setdefault(step, []).append(
        f"[{time.strftime('%H:%M:%S')}] {message}"
    )


def _render_logs(step: str) -> None:
    logs = st.session_state.step_logs.get(step, [])
    if logs:
        st.markdown("#### 📝 Execution Log")
        st.markdown(
            f'<div class="log-line">{"<br>".join(logs)}</div>',
            unsafe_allow_html=True,
        )


def _download_button(data: bytes | str, file_name: str, mime: str, label: str = "Download", key: str | None = None) -> None:
    if isinstance(data, str):
        data = data.encode("utf-8")
    kwargs = {"label": label, "data": data, "file_name": file_name, "mime": mime}
    if key:
        kwargs["key"] = key
    st.download_button(**kwargs)


def _download_json(data: dict[str, Any], file_name: str, label: str = "📥 Download JSON", key: str | None = None) -> None:
    _download_button(
        json.dumps(data, ensure_ascii=False, indent=2),
        file_name,
        "application/json",
        label,
        key=key or f"dl_{file_name}",
    )


# ---------------------------------------------------------------------------
# Pipeline Map Visualisation
# ---------------------------------------------------------------------------
PIPELINE_STEPS = [
    ("setup", "📁 Setup", "Upload files & configure"),
    ("load_case", "📂 Load Case", "Detect topology"),
    ("base_analysis", "🏛️ Base Analysis", "Regex/OOXML structure"),
    ("amendment_analysis", "📝 Amendment Analysis", "LLM + regex extraction"),
    ("analysis_gate", "🛡️ Analysis Gate", "Coverage & repair check"),
    ("skeleton", "🦴 Skeleton Build", "Service tables"),
    ("resolution", "🎯 Resolution", "Resolve intents to operations"),
    ("manual_review", "🔍 Manual Review", "Safe vs blocked split"),
    ("edit", "✏️ Edit", "Apply OOXML mutations"),
    ("markers", "📌 Markers", "Insert revision markers"),
    ("reanalyze", "🔄 Reanalyze", "Final structure check"),
    ("checklist", "📋 Checklist", "Build validation list"),
    ("validation", "⚖️ Validation", "LLM Judge + deterministic checks"),
]

STEP_GROUPS = [
    ["setup"],
    ["load_case"],
    ["base_analysis", "amendment_analysis"],
    ["analysis_gate"],
    ["skeleton"],
    ["resolution"],
    ["manual_review"],
    ["edit"],
    ["markers"],
    ["reanalyze"],
    ["checklist"],
    ["validation"],
]


def _step_css_class(step: str) -> str:
    status = st.session_state.step_status.get(step, "pending")
    return f"step-node step-{status}"


def _step_display_name(step: str) -> str:
    for key, title, _ in PIPELINE_STEPS:
        if key == step:
            return title
    return step


def render_pipeline_map() -> None:
    st.markdown("### 🗺️ Pipeline Map")
    rows_html = ""
    for group in STEP_GROUPS:
        cells = ""
        for step in group:
            _, title, label = next((s for s in PIPELINE_STEPS if s[0] == step), (step, step, ""))
            css = _step_css_class(step)
            cells += f'<div class="{css}">{title}<div class="step-label">{label}</div></div>'
        rows_html += f'<div class="pipeline-row">{cells}</div>'
        if group != STEP_GROUPS[-1]:
            rows_html += '<div class="arrow-down">▼</div>'
    st.markdown(
        f'<div class="pipeline-map">{rows_html}</div>',
        unsafe_allow_html=True,
    )


# ---------------------------------------------------------------------------
# Step Runners
# ---------------------------------------------------------------------------
def _get_agents() -> dict[str, Any]:
    if "amendment_analyzer" not in st.session_state.agents:
        config = load_models_config(st.session_state.models_config_path)
        st.session_state.agents = {
            "config": config,
            "amendment_analyzer": AmendmentAnalyzer(config),
            "base_analyzer": BaseAnalyzer(),
            "skeleton_builder": SkeletonBuilder(),
            "resolver": PipelineResolver(config),
            "editor": EditorV2(),
            "checklist_builder": ValidationChecklistBuilder(),
            "validator": StrictJudgeValidator(config),
            "marker_inserter": RevisionMarkerInserter(),
        }
    return st.session_state.agents


def run_step_setup() -> bool:
    step = "setup"
    _set_step_status(step, "running")
    _add_log(step, "Initializing workspace...")
    try:
        base_file = st.session_state.base_file
        amendment_file = st.session_state.amendment_file
        tmpdir = st.session_state.workspace_dir
        workspace = Path(tmpdir) / "workspace"
        workspace.mkdir(parents=True, exist_ok=True)
        case_id = "web_ui_case"
        case_dir = workspace / "redactions" / case_id
        case_dir.mkdir(parents=True, exist_ok=True)

        base_path = case_dir / base_file.name
        amend_path = case_dir / amendment_file.name
        base_path.write_bytes(base_file.getvalue())
        amend_path.write_bytes(amendment_file.getvalue())

        models_config = get_default_models_config()
        # apply sidebar overrides if present (stored in session)
        models_config["runtime"]["base_url"] = st.session_state.get("base_url", models_config["runtime"]["base_url"])
        models_config["runtime"]["api_key"] = st.session_state.get("api_key", models_config["runtime"]["api_key"])
        models_config["runtime"]["temperature"] = st.session_state.get("temperature", models_config["runtime"]["temperature"])
        models_config["runtime"]["timeout"] = st.session_state.get("timeout", models_config["runtime"]["timeout"])
        models_config["runtime"]["semantic_ranking_enabled"] = st.session_state.get("semantic_ranking", False)
        model_name = st.session_state.get("model_name", models_config["models"]["default"])
        for role in ["default", "analyst", "resolver_disambiguation", "validator"]:
            models_config["models"][role] = model_name

        models_path = Path(tmpdir) / "models.json"
        models_path.write_text(json.dumps(models_config, ensure_ascii=False, indent=2), encoding="utf-8")

        st.session_state.workspace = workspace
        st.session_state.models_config_path = models_path
        st.session_state.case_id = case_id
        st.session_state.base_path = base_path
        st.session_state.amendment_path = amend_path
        st.session_state.checklist = PipelineChecklist(case_id)

        _add_log(step, f"Workspace: {workspace}")
        _add_log(step, f"Base: {base_path.name}")
        _add_log(step, f"Amendment: {amend_path.name}")
        _add_log(step, "Setup complete.")
        _set_step_status(step, "success")
        return True
    except Exception as exc:
        _add_log(step, f"ERROR: {exc}")
        _set_step_status(step, "error")
        st.session_state.pipeline_error = exc
        return False


def run_step_load_case() -> bool:
    step = "load_case"
    _set_step_status(step, "running")
    _add_log(step, "Running case_loader...")
    try:
        case = load_case(st.session_state.workspace, st.session_state.case_id)
        st.session_state.case = case
        topology = case.get("case_topology", "unknown")
        st.session_state.base_doc = case.get("base_doc")
        st.session_state.amendment_paths = case["amendment_docs"]
        _record_initial_checks(
            st.session_state.checklist,
            case_topology=topology,
            amendment_paths=st.session_state.amendment_paths,
            base_docs=[st.session_state.base_doc] if st.session_state.base_doc else [],
        )
        _add_log(step, f"Topology: {topology}")
        _add_log(step, f"Amendments: {len(st.session_state.amendment_paths)}")
        _add_log(step, f"Base doc: {st.session_state.base_doc.name if st.session_state.base_doc else 'none'}")
        _set_step_status(step, "success")
        return True
    except Exception as exc:
        _add_log(step, f"ERROR: {exc}")
        _set_step_status(step, "error")
        st.session_state.pipeline_error = exc
        return False


def run_step_base_analysis() -> bool:
    step = "base_analysis"
    _set_step_status(step, "running")
    _add_log(step, "Starting BaseAnalyzer...")
    try:
        agents = _get_agents()
        base_doc = st.session_state.base_doc
        analysis = agents["base_analyzer"].analyze(base_doc)
        st.session_state.base_analysis = analysis
        _add_log(step, f"Complexity: {analysis.complexity}")
        _add_log(step, f"Header blocks: {len(analysis.header_blocks)}")
        for hb in analysis.header_blocks:
            _add_log(step, f"  [{hb.scope}] {hb.header_id} -> after para {hb.end_paragraph_index}")
        _set_step_status(step, "success")
        return True
    except Exception as exc:
        _add_log(step, f"ERROR: {exc}")
        _set_step_status(step, "error")
        st.session_state.pipeline_error = exc
        return False


def run_step_amendment_analysis() -> bool:
    step = "amendment_analysis"
    _set_step_status(step, "running")
    _add_log(step, "Starting AmendmentAnalyzer (LLM + regex)...")
    try:
        agents = _get_agents()
        paths = st.session_state.amendment_paths
        started = time.perf_counter()
        analyses = agents["amendment_analyzer"].analyze_many(paths)
        elapsed = time.perf_counter() - started
        st.session_state.amendment_analyses = analyses
        total_intents = sum(len(a.intents) for a in analyses)
        _add_log(step, f"Documents analyzed: {len(analyses)}")
        _add_log(step, f"Total intents: {total_intents}")
        _add_log(step, f"Elapsed: {elapsed:.1f}s")
        for a in analyses:
            _add_log(step, f"  [{a.metadata.document_label}] {len(a.intents)} intents, complexity={a.metadata.complexity}")
        _set_step_status(step, "success")
        return True
    except Exception as exc:
        _add_log(step, f"ERROR: {exc}")
        _set_step_status(step, "error")
        st.session_state.pipeline_error = exc
        return False


def run_step_analysis_gate() -> bool:
    step = "analysis_gate"
    _set_step_status(step, "running")
    _add_log(step, "Checking analysis gate...")
    try:
        agents = _get_agents()
        analyses = st.session_state.amendment_analyses
        base_analysis_dict = {str(st.session_state.base_doc): st.session_state.base_analysis}
        paths = st.session_state.amendment_paths

        # repair once if needed
        analyses = _repair_analysis_once(
            st.session_state.checklist,
            case_id=st.session_state.case_id,
            amendment_analyzer=agents["amendment_analyzer"],
            amendment_paths=paths,
            amendment_analyses=analyses,
            base_analyses=base_analysis_dict,
        )
        st.session_state.amendment_analyses = analyses

        ok, details = _analysis_gate_details(analyses, base_analysis_dict, paths)
        st.session_state.analysis_gate_details = details
        _add_log(step, f"Gate OK: {ok}")
        _add_log(step, f"Estimated directives vs intents: {json.dumps({p.name: d['intent_count'] for p, d in zip(paths, details['coverage'])}, ensure_ascii=False)}")
        if not ok:
            _add_log(step, "WARNING: Gate failed — some directives not covered.")
            st.session_state.blocked_operations = _analysis_blocked_operations(details)
            _set_step_status(step, "warning")
            # We don't stop here in stepped mode so user can inspect
            return True
        _set_step_status(step, "success")
        return True
    except Exception as exc:
        _add_log(step, f"ERROR: {exc}")
        _set_step_status(step, "error")
        st.session_state.pipeline_error = exc
        return False


def run_step_skeleton() -> bool:
    step = "skeleton"
    _set_step_status(step, "running")
    _add_log(step, "Building skeleton + service tables...")
    try:
        agents = _get_agents()
        artifacts_dir = st.session_state.workspace / "artifacts" / st.session_state.case_id
        artifacts_dir.mkdir(parents=True, exist_ok=True)
        skeleton_doc = artifacts_dir / "working_skeleton.docx"
        specs = agents["skeleton_builder"].build(
            st.session_state.base_doc,
            st.session_state.base_analysis,
            st.session_state.amendment_analyses,
            skeleton_doc,
        )
        st.session_state.skeleton_doc = skeleton_doc
        st.session_state.service_table_specs = specs
        st.session_state.current_doc = skeleton_doc
        _add_log(step, f"Skeleton: {skeleton_doc.name}")
        _add_log(step, f"Service table specs: {len(specs)}")
        for sp in specs:
            _add_log(step, f"  [{sp.scope}] labels={len(sp.document_labels)} after_para={sp.insert_after_paragraph_index}")
        _set_step_status(step, "success")
        return True
    except Exception as exc:
        _add_log(step, f"ERROR: {exc}")
        _set_step_status(step, "error")
        st.session_state.pipeline_error = exc
        return False


def run_step_resolution() -> bool:
    step = "resolution"
    _set_step_status(step, "running")
    _add_log(step, "Resolving intents to operations...")
    try:
        agents = _get_agents()
        # For simplicity we handle the first amendment in stepped mode (standard_single)
        amendment_analysis = st.session_state.amendment_analyses[0]
        current_doc = st.session_state.current_doc
        resolution = agents["resolver"].resolve(current_doc, amendment_analysis.intents)
        resolved_ops = resolution["resolved_operations"]
        resolved_count = sum(1 for o in resolved_ops if o.status == "resolved")
        ambiguous_count = sum(1 for o in resolved_ops if o.status == "ambiguous")
        unsupported_count = sum(1 for o in resolved_ops if o.status == "unsupported")

        st.session_state.resolution = resolution
        _add_log(step, f"Total operations: {len(resolved_ops)}")
        _add_log(step, f"Resolved: {resolved_count}, Ambiguous: {ambiguous_count}, Unsupported: {unsupported_count}")

        # gate (non-raising)
        gate_ok = _enforce_resolution_gate(
            st.session_state.checklist,
            case_id=st.session_state.case_id,
            base_doc=st.session_state.base_doc,
            amendment_analysis=amendment_analysis,
            amendment_index=1,
            resolved_operations=resolved_ops,
            pass_name="initial",
            raise_on_fail=False,
        )
        if not gate_ok:
            _add_log(step, "Resolution gate failed — attempting repair...")
            resolution = agents["resolver"].resolve(current_doc, amendment_analysis.intents, repair=True)
            resolved_ops = resolution["resolved_operations"]
            st.session_state.resolution = resolution
            resolved_count = sum(1 for o in resolved_ops if o.status == "resolved")
            _add_log(step, f"After repair — Resolved: {resolved_count}")

        _set_step_status(step, "success")
        return True
    except Exception as exc:
        _add_log(step, f"ERROR: {exc}")
        _set_step_status(step, "error")
        st.session_state.pipeline_error = exc
        return False


def run_step_manual_review() -> bool:
    step = "manual_review"
    _set_step_status(step, "running")
    _add_log(step, "Splitting operations for manual review...")
    try:
        resolved_ops = st.session_state.resolution["resolved_operations"]
        split = split_operations_for_manual_review(resolved_ops)
        st.session_state.split = split
        safe = len(split.safe_to_apply)
        blocked = len(split.blocked_operations)
        _add_log(step, f"Safe to apply: {safe}")
        _add_log(step, f"Blocked: {blocked}")
        if blocked:
            for b in split.blocked_operations:
                _add_log(step, f"  BLOCKED [{b.get('operation_kind', '?')}] {b.get('reason', '')}")
            _set_step_status(step, "warning")
        else:
            _set_step_status(step, "success")
        return True
    except Exception as exc:
        _add_log(step, f"ERROR: {exc}")
        _set_step_status(step, "error")
        st.session_state.pipeline_error = exc
        return False


def run_step_edit() -> bool:
    step = "edit"
    _set_step_status(step, "running")
    _add_log(step, "Applying edits via EditorV2...")
    try:
        agents = _get_agents()
        split = st.session_state.split
        if not split.safe_to_apply:
            _add_log(step, "No safe operations to apply — skipping edit.")
            _set_step_status(step, "warning")
            return True
        artifacts_dir = st.session_state.workspace / "artifacts" / st.session_state.case_id
        step_output = artifacts_dir / "working_step_1.docx"
        edit_result = agents["editor"].edit(st.session_state.current_doc, step_output, split.safe_to_apply)
        st.session_state.edit_result = edit_result
        st.session_state.current_doc = step_output
        st.session_state.all_operations.extend(split.safe_to_apply)
        st.session_state.all_statuses.extend(edit_result.get("statuses", []))
        _add_log(step, f"Output: {step_output.name}")
        for status_line in edit_result.get("statuses", [])[:8]:
            _add_log(step, f"  {status_line}")
        if len(edit_result.get("statuses", [])) > 8:
            _add_log(step, f"  ... and {len(edit_result['statuses']) - 8} more")
        drift = edit_result.get("drift_events", [])
        if drift:
            _add_log(step, f"Drift events: {len(drift)}")
        _set_step_status(step, "success")
        return True
    except Exception as exc:
        _add_log(step, f"ERROR: {exc}")
        _set_step_status(step, "error")
        st.session_state.pipeline_error = exc
        return False


def run_step_markers() -> bool:
    step = "markers"
    _set_step_status(step, "running")
    _add_log(step, "Inserting revision markers...")
    try:
        agents = _get_agents()
        split = st.session_state.split
        if not split.safe_to_apply:
            _add_log(step, "No safe operations — skipping markers.")
            _set_step_status(step, "warning")
            return True
        marker_result = agents["marker_inserter"].insert_markers(st.session_state.current_doc, split.safe_to_apply)
        st.session_state.marker_result = marker_result
        _add_log(step, f"Markers inserted: {len(marker_result)}")
        for m in marker_result[:8]:
            _add_log(step, f"  {m}")
        if len(marker_result) > 8:
            _add_log(step, f"  ... and {len(marker_result) - 8} more")
        _set_step_status(step, "success")
        return True
    except Exception as exc:
        _add_log(step, f"ERROR: {exc}")
        _set_step_status(step, "error")
        st.session_state.pipeline_error = exc
        return False


def run_step_reanalyze() -> bool:
    step = "reanalyze"
    _set_step_status(step, "running")
    _add_log(step, "Reanalyzing final document + final service tables...")
    try:
        agents = _get_agents()
        final_base_analysis = agents["base_analyzer"].analyze(st.session_state.current_doc)
        final_specs = build_service_table_specs(final_base_analysis, st.session_state.amendment_analyses)
        doc = Document(st.session_state.current_doc)
        doc.save(st.session_state.current_doc)
        st.session_state.final_base_analysis = final_base_analysis
        st.session_state.final_specs = final_specs
        _add_log(step, f"Final header blocks: {len(final_base_analysis.header_blocks)}")
        _add_log(step, f"Final service tables: {len(final_specs)}")
        _set_step_status(step, "success")
        return True
    except Exception as exc:
        _add_log(step, f"ERROR: {exc}")
        _set_step_status(step, "error")
        st.session_state.pipeline_error = exc
        return False


def run_step_checklist() -> bool:
    step = "checklist"
    _set_step_status(step, "running")
    _add_log(step, "Building validation checklist...")
    try:
        agents = _get_agents()
        checklist = agents["checklist_builder"].build(
            base_analysis=st.session_state.final_base_analysis,
            amendment_analyses=st.session_state.amendment_analyses,
            service_table_specs=st.session_state.final_specs,
            resolved_operations=st.session_state.all_operations,
            runtime_checks=st.session_state.checklist.items(),
        )
        st.session_state.validation_checklist = checklist
        _add_log(step, f"Checks created: {len(checklist.checks)}")
        _set_step_status(step, "success")
        return True
    except Exception as exc:
        _add_log(step, f"ERROR: {exc}")
        _set_step_status(step, "error")
        st.session_state.pipeline_error = exc
        return False


def run_step_validation() -> bool:
    step = "validation"
    _set_step_status(step, "running")
    _add_log(step, "Running StrictJudgeValidator...")
    try:
        agents = _get_agents()
        operation_summary = {
            "total": len(st.session_state.all_operations),
            "resolved": sum(1 for o in st.session_state.all_operations if o.status == "resolved"),
            "ambiguous": sum(1 for o in st.session_state.all_operations if o.status == "ambiguous"),
            "unsupported": sum(1 for o in st.session_state.all_operations if o.status == "unsupported"),
        }
        _add_log(step, f"Operation summary: {operation_summary}")
        validation = agents["validator"].validate(
            output_doc=st.session_state.current_doc,
            checklist=st.session_state.validation_checklist,
            amendment_analyses=st.session_state.amendment_analyses,
            base_analysis=st.session_state.final_base_analysis,
            operation_statuses=st.session_state.all_statuses,
            operation_summary=operation_summary,
        )
        st.session_state.validation = validation
        vd = validation.to_dict()
        _add_log(step, f"Structural OK: {vd['structural_ok']}")
        _add_log(step, f"Judge OK: {vd['judge_ok']}")
        _add_log(step, f"Is Valid: {vd['is_valid']}")
        if vd.get("judge_summary"):
            _add_log(step, f"Summary: {vd['judge_summary']}")
        if vd.get("judge_failures"):
            for f in vd["judge_failures"][:5]:
                _add_log(step, f"  FAIL: {f}")
        _set_step_status(step, "success" if vd["is_valid"] else "warning")
        return True
    except Exception as exc:
        _add_log(step, f"ERROR: {exc}")
        _set_step_status(step, "error")
        st.session_state.pipeline_error = exc
        return False


# ---------------------------------------------------------------------------
# Step UI Renderers (results after execution)
# ---------------------------------------------------------------------------
def _build_full_result() -> dict[str, Any]:
    """Собирает полный result dict из session_state (как в CLI)."""
    result: dict[str, Any] = {
        "case_id": st.session_state.get("case_id"),
        "case_topology": st.session_state.case.get("case_topology") if st.session_state.get("case") else None,
        "workspace_root": str(st.session_state.workspace) if st.session_state.get("workspace") else None,
    }
    if st.session_state.get("base_doc"):
        result["base_doc"] = str(st.session_state.base_doc)
    if st.session_state.get("base_analysis"):
        result["base_analysis"] = st.session_state.base_analysis.to_dict()
    if st.session_state.get("amendment_analyses"):
        result["amendments"] = [a.to_dict() for a in st.session_state.amendment_analyses]
    if st.session_state.get("service_table_specs"):
        result["service_table_specs"] = [s.to_dict() for s in st.session_state.service_table_specs]
    steps_data: list[dict[str, Any]] = []
    step: dict[str, Any] = {}
    if st.session_state.get("amendment_analyses"):
        step["amendment"] = st.session_state.amendment_analyses[0].to_dict()
    if st.session_state.get("resolution"):
        step["resolution"] = {
            "resolved_operations": [o.to_dict() for o in st.session_state.resolution.get("resolved_operations", [])],
            "debug_candidates": st.session_state.resolution.get("debug_candidates"),
        }
        if st.session_state.get("split"):
            step["resolution"]["safe_to_apply"] = [o.to_dict() for o in st.session_state.split.safe_to_apply]
            step["resolution"]["blocked_operations"] = st.session_state.split.blocked_operations
    if st.session_state.get("edit_result"):
        step["edit"] = st.session_state.edit_result
    if st.session_state.get("marker_result"):
        step["revision_markers"] = st.session_state.marker_result
    if st.session_state.get("current_doc"):
        step["output_doc"] = str(st.session_state.current_doc)
    if step:
        steps_data.append(step)
    if steps_data:
        result["steps"] = steps_data
    if st.session_state.get("current_doc"):
        result["output_doc"] = str(st.session_state.current_doc)
    if st.session_state.get("final_base_analysis"):
        result["base_analysis"] = st.session_state.final_base_analysis.to_dict()
    if st.session_state.get("final_specs"):
        result["service_table_specs"] = [s.to_dict() for s in st.session_state.final_specs]
    if st.session_state.get("validation_checklist"):
        result["validation_checklist"] = st.session_state.validation_checklist.to_dict()
    if st.session_state.get("validation"):
        vd = st.session_state.validation.to_dict()
        result["validation"] = vd
        result["status"] = vd.get("status", "unknown")
        result["manual_review_required"] = vd.get("manual_review_required", False)
    if st.session_state.get("blocked_operations"):
        result["blocked_operations"] = st.session_state.blocked_operations
    safe_count = len(st.session_state.get("all_operations", []))
    result["safe_operations_applied"] = safe_count
    return result


def _render_step_base_analysis() -> None:
    analysis = st.session_state.base_analysis
    if not analysis:
        return
    st.markdown("**Complexity:** `" + analysis.complexity + "`")
    st.markdown(f"**Header blocks:** {len(analysis.header_blocks)}")
    for hb in analysis.header_blocks:
        with st.expander(f"[{hb.scope}] {hb.header_id}"):
            st.json(hb.to_dict())
    _download_json({"base_analysis": analysis.to_dict()}, "base_analysis.json", key="dl_base_analysis")


def _render_step_amendment_analysis() -> None:
    analyses = st.session_state.amendment_analyses
    if not analyses:
        return
    for a in analyses:
        meta = a.metadata
        with st.expander(f"{meta.document_label} ({len(a.intents)} intents)"):
            st.markdown(f"**Complexity:** `{meta.complexity}`")
            st.markdown(f"**Document number:** {meta.document_number}")
            st.markdown(f"**Date:** {meta.document_date_iso}")
            st.markdown("**Intents:**")
            for intent in a.intents:
                with st.container():
                    c1, c2 = st.columns([1, 4])
                    with c1:
                        st.code(intent.operation_kind)
                    with c2:
                        excerpt = intent.source_excerpt[:120] if intent.source_excerpt else ""
                        st.write(f"`{intent.change_id}` — {excerpt}...")
            if st.checkbox("Show raw JSON", key=f"raw_{meta.document_label}"):
                st.json(a.to_dict())
    _download_json({"amendments": [a.to_dict() for a in analyses]}, "amendment_analysis.json", key="dl_amendment_analysis")


def _render_step_analysis_gate() -> None:
    details = st.session_state.get("analysis_gate_details", {})
    st.markdown("**Coverage Details:**")
    for item in details.get("coverage", []):
        ok = item.get("ok", False)
        icon = "✅" if ok else "⚠️"
        st.write(f"{icon} `{Path(item['source_path']).name}` — intents={item['intent_count']} vs estimated={item['estimated_directives']}")
    if details.get("coverage_failed"):
        st.warning("Some directives were not converted to intents!")
        for fail in details["coverage_failed"]:
            st.write(f"- Missing in `{Path(fail['source_path']).name}`")
    if st.session_state.blocked_operations:
        st.markdown(f"**Blocked operations:** {len(st.session_state.blocked_operations)}")
    _download_json({"analysis_gate": details, "blocked_operations": st.session_state.blocked_operations}, "analysis_gate.json", key="dl_analysis_gate")


def _render_step_skeleton() -> None:
    specs = st.session_state.service_table_specs
    st.markdown(f"**Skeleton doc:** `{st.session_state.skeleton_doc.name if st.session_state.skeleton_doc else '—'}`")
    st.markdown(f"**Service table specs:** {len(specs)}")
    for sp in specs:
        st.write(f"- `[{sp.scope}]` after_para={sp.insert_after_paragraph_index} labels={len(sp.document_labels)}")
    if st.session_state.skeleton_doc and st.session_state.skeleton_doc.exists():
        with st.expander("Preview skeleton text"):
            st.text(extract_docx_preview(st.session_state.skeleton_doc, max_paras=20))
    _download_json(
        {"service_table_specs": [s.to_dict() for s in specs], "skeleton_doc": str(st.session_state.skeleton_doc) if st.session_state.skeleton_doc else None},
        "skeleton.json",
        key="dl_skeleton",
    )


def _render_step_resolution() -> None:
    resolution = st.session_state.resolution
    if not resolution:
        return
    ops = resolution["resolved_operations"]
    candidates = resolution.get("debug_candidates", {})
    statuses = {"resolved": 0, "ambiguous": 0, "unsupported": 0}
    for o in ops:
        statuses[o.status] = statuses.get(o.status, 0) + 1
    c1, c2, c3 = st.columns(3)
    c1.metric("Resolved", statuses.get("resolved", 0))
    c2.metric("Ambiguous", statuses.get("ambiguous", 0))
    c3.metric("Unsupported", statuses.get("unsupported", 0))
    with st.expander(f"Operations ({len(ops)} total)"):
        for op in ops:
            with st.container():
                col_badge, col_body = st.columns([1, 4])
                with col_badge:
                    color = {"resolved": "green", "ambiguous": "orange", "unsupported": "red"}.get(op.status, "gray")
                    st.markdown(f":{color}-badge[{op.status}]")
                with col_body:
                    st.write(f"`{op.operation_id}` → `{op.operation_kind}`")
                    if op.ambiguity_reason:
                        st.caption(f"Reason: {op.ambiguity_reason}")
    if candidates:
        with st.expander("Debug candidates"):
            st.json({k: [c.to_dict() for c in v] for k, v in candidates.items()})
    _download_json(
        {
            "resolved_operations": [o.to_dict() for o in ops],
            "debug_candidates": {k: [c.to_dict() for c in v] for k, v in candidates.items()},
        },
        "resolution.json",
        key="dl_resolution",
    )


def _render_step_manual_review() -> None:
    split = st.session_state.split
    if not split:
        return
    c1, c2 = st.columns(2)
    with c1:
        st.metric("Safe to apply", len(split.safe_to_apply))
    with c2:
        st.metric("Blocked", len(split.blocked_operations))
    if split.safe_to_apply:
        with st.expander("Safe operations"):
            for op in split.safe_to_apply:
                st.write(f"- `{op.operation_id}` {op.operation_kind}")
    if split.blocked_operations:
        with st.expander("Blocked operations"):
            for b in split.blocked_operations:
                st.error(f"`{b.get('operation_kind', '?')}` — {b.get('reason', '')}")
    _download_json(
        {
            "safe_to_apply": [o.to_dict() for o in split.safe_to_apply],
            "blocked_operations": split.blocked_operations,
        },
        "manual_review.json",
        key="dl_manual_review",
    )


def _render_step_edit() -> None:
    edit = st.session_state.edit_result
    if not edit:
        st.info("No edit performed (no safe operations).")
        return
    st.markdown(f"**Statuses:** {len(edit.get('statuses', []))}")
    for s in edit.get("statuses", [])[:10]:
        st.write(f"- {s}")
    if len(edit.get("statuses", [])) > 10:
        st.caption(f"... and {len(edit['statuses']) - 10} more")
    drift = edit.get("drift_events", [])
    if drift:
        st.markdown(f"**Drift events:** {len(drift)}")
        for d in drift:
            st.json(d)
    if st.session_state.current_doc and st.session_state.current_doc.exists():
        with st.expander("Preview edited document"):
            st.text(extract_docx_preview(st.session_state.current_doc, max_paras=20))


def _render_step_markers() -> None:
    markers = st.session_state.marker_result
    if not markers:
        st.info("No markers inserted.")
        return
    st.markdown(f"**Markers:** {len(markers)}")
    for m in markers[:15]:
        st.write(f"- {m}")
    if len(markers) > 15:
        st.caption(f"... and {len(markers) - 15} more")
    _download_json({"revision_markers": markers}, "revision_markers.json", key="dl_markers")


def _render_step_reanalyze() -> None:
    analysis = st.session_state.final_base_analysis
    if analysis:
        st.markdown(f"**Final header blocks:** {len(analysis.header_blocks)}")
        st.markdown(f"**Final service table specs:** {len(st.session_state.final_specs)}")
        _download_json(
            {
                "final_base_analysis": analysis.to_dict(),
                "final_service_table_specs": [s.to_dict() for s in st.session_state.final_specs],
            },
            "reanalyze.json",
            key="dl_reanalyze",
        )


def _render_step_checklist() -> None:
    cl = st.session_state.validation_checklist
    if cl:
        st.markdown(f"**Checks:** {len(cl.checks)}")
        with st.expander("View checks"):
            for ch in cl.checks:
                icon = "✅" if ch.get("ok") else "❌"
                st.write(f"{icon} `{ch.get('stage', '?')}` / `{ch.get('check_id', '?')}`")
        _download_json(cl.to_dict(), "validation_checklist.json", key="dl_checklist")


def _render_step_validation() -> None:
    val = st.session_state.validation
    if not val:
        return
    vd = val.to_dict()
    c1, c2, c3 = st.columns(3)
    with c1:
        st.metric("Structural OK", "✅" if vd["structural_ok"] else "❌")
    with c2:
        st.metric("Judge OK", "✅" if vd["judge_ok"] else "❌")
    with c3:
        color = "normal" if vd["is_valid"] else "inverse"
        st.metric("Is Valid", "✅ YES" if vd["is_valid"] else "❌ NO")
    if vd.get("judge_summary"):
        st.markdown(f"**Summary:** {vd['judge_summary']}")
    if vd.get("judge_failures"):
        st.markdown("**Failures:**")
        for f in vd["judge_failures"]:
            st.error(f)
    if vd.get("intent_results"):
        with st.expander("Intent check details"):
            for ir in vd["intent_results"]:
                ok = ir.get("check_ok", False)
                st.write(f"{'✅' if ok else '❌'} `{ir.get('intent_id', '?')}` ({ir.get('check_type', '?')})")
    _download_json(vd, "validation.json", key="dl_validation_step")


STEP_RUNNERS = {
    "setup": run_step_setup,
    "load_case": run_step_load_case,
    "base_analysis": run_step_base_analysis,
    "amendment_analysis": run_step_amendment_analysis,
    "analysis_gate": run_step_analysis_gate,
    "skeleton": run_step_skeleton,
    "resolution": run_step_resolution,
    "manual_review": run_step_manual_review,
    "edit": run_step_edit,
    "markers": run_step_markers,
    "reanalyze": run_step_reanalyze,
    "checklist": run_step_checklist,
    "validation": run_step_validation,
}

STEP_RENDERERS = {
    "base_analysis": _render_step_base_analysis,
    "amendment_analysis": _render_step_amendment_analysis,
    "analysis_gate": _render_step_analysis_gate,
    "skeleton": _render_step_skeleton,
    "resolution": _render_step_resolution,
    "manual_review": _render_step_manual_review,
    "edit": _render_step_edit,
    "markers": _render_step_markers,
    "reanalyze": _render_step_reanalyze,
    "checklist": _render_step_checklist,
    "validation": _render_step_validation,
}

# ---------------------------------------------------------------------------
# Step Descriptions (what happens + which code is responsible)
# ---------------------------------------------------------------------------
STEP_DESCRIPTIONS: dict[str, str] = {
    "setup": """
**Назначение:** подготовка временного workspace, копирование загруженных .docx, формирование `models.json`.

**Ключевой код:**
- `run_step_setup()` в `web_ui_stepped.py` — копирует файлы в `workspace/redactions/<case_id>/`
- Формирует `models_config` из настроек sidebar (`get_default_models_config()`)
- Пишет `models.json` на диск

**Почему так:** Pipeline ожидает строгую файловую структуру (`redactions/<case-id>/*.docx`), чтобы `case_loader` смог определить топологию по именам файлов.
""",
    "load_case": """
**Назначение:** авто-определение топологии кейса по содержимому директории.

**Ключевой код:**
- `redacta/case_loader.py` → `load_case(workspace, case_id)`
- `_real_docx_files()` — фильтрует временные `~$*` и `updated*`
- `_is_amendment_file()` — regex: имя начинается с `изм` или `изменение` → amendment; всё остальное → base

**Логика определения топологии:**
| Условие | Топология |
|---|---|
| 1 amend + 1 base | `standard_single` |
| ≥2 amend + 1 base | `special_multi_amendment_single_base` |
| 1 amend + ≥2 base | `special_single_amendment_multi_base` |

**Почему regex по имени файла, а не LLM:** имена файлов — соглашение по проекту. LLM здесь избыточен и медленнее.
""",
    "base_analysis": """
**Назначение:** построить структурный «скелет» базового документа: `header_blocks` (преамбула + приложения) и `complexity`.

**Ключевой код:**
- `redacta/base_analyzer.py` → `BaseAnalyzer.analyze(base_doc)`
- `_find_top_header_end()` — ищет границу шапки по boundary-фразам (`приказываю:`, `в соответствии`, `^\\d+\\.\\s`)
- `_extract_appendix_number()` — regex `\\b(\\d+)\\b`
- `_find_format_appendix_headers()` — ищет паттерн «правый блок → центрированный блок» через `WD_ALIGN_PARAGRAPH`
- `document_classifier.py` → `classify_base_complexity()`: `media_heavy > table_heavy > plain`

**Почему regex/OOXML, а не LLM:** структура шапки детерминирована (обязательные слова-границы, стилевые признаки). LLM дольше, дороже и может «галлюцинировать» приложения.
""",
    "amendment_analysis": """
**Назначение:** извлечь `ChangeIntent` из документа изменений — семантический парсинг директив.

**Ключевой код:**
- `redacta/amendment_analyzer.py` → `AmendmentAnalyzer.analyze_many(paths)`
- Внутри композиция:
  1. `DeterministicIntentExtractor.extract(lines)` — pure regex (`исключить|заменить|утратившим силу`) → 100% precision, confidence 1.0
  2. `AmendmentLLMAnalyzer.analyze()` — LLM-вызов (`system.txt` + `user_template.txt`, max_tokens=2000)
  3. `_fallback_extract(lines)` — доп. regex для 7 видов операций
  4. `_merge_with_fallback()` — объединение LLM + regex
  5. `_normalize_intents()` — 15+ regex-проходов «дотачивания» адресации

**Почему двухуровневый (LLM + regex):**
- Regex покрывает канонические формулировки (`пункт N признать утратившим силу`) — мгновенно и точно.
- LLM парсит «творческий» хвост: «в подпункте «в» пункта 2 абзац третий изложить в редакции…» — семантическая адресация.
- Post-processing (`_normalize_intents`) дешевле, чем заставлять LLM возвращать идеально структурированный JSON.
""",
    "analysis_gate": """
**Назначение:** гарантировать, что ни одна директива не пропущена перед редактированием.

**Ключевой код:**
- `run_case.py` → `_extract_amendment_directives(amendment_doc)`
  - Читает параграфы, входит в «режим директив» после первого `^\\d+\\.\\s` (пропускает преамбулу)
  - Оставляет строки с маркерами `исключить|заменить|изложить|дополнить|признать.*утрат`
  - `_estimate_action_units()` — считает quoted-фразы (`слова "А", "Б" исключить` → 3 атомарные единицы)
- `_analysis_gate_details()` — сравнивает `intent_count >= expected_count`
- `_repair_analysis_once()` — если gate не пройден:
  - Вызывает `AmendmentAnalyzer.repair_analysis(path, prev_analysis, directives)` с `repair_user_template.txt`
  - Контекст 240 строк, max_tokens=3000

**Почему gate важен:** один пропущенный intent = одна неприменённая правка = некорректный итоговый документ.
""",
    "skeleton": """
**Назначение:** создать рабочую копию базового документа и вставить служебные таблицы (стиль «Консультант-плюс»).

**Ключевой код:**
- `redacta/skeleton_builder.py` → `SkeletonBuilder.build(base_doc, base_analysis, amendments, output_doc)`
- `redacta/service_tables.py`:
  - `build_service_table_specs()` — для каждого `header_block` формирует `ServiceTableSpec`
  - `insert_service_tables()` — сортирует specs по `insert_after_paragraph_index` в обратном порядке (чтобы индексы не сдвигались)
  - `_insert_service_table_ooxml()` — создаёт `<w:tbl>` с фиксированной grid `[60, 113, 9921, 113]`, фоном `CED3F1`/`F4F3F8`, цветом текста `392C69`

**Почему OOXML напрямую:** python-docx high-level API не позволяет контролировать `tcPr`/`tblBorders`/`tcMar` на нужном уровне. High-level API ставит дефолтные стили, портящие визуальный паттерн.
""",
    "resolution": """
**Назначение:** сопоставить каждый `ChangeIntent` с конкретным местом в базовом документе (`ResolvedOperation`).

**Ключевой код:**
- `redacta/resolver.py` — тонкий wrapper
- `redacta/resolver_v2.py` → `ResolverV2.resolve(base_doc, intents, mode="anchor_id")`
  1. `read_paragraph_records()` — список `ParagraphRecord(absolute_index, text)`
  2. Switch по `operation_kind` → `_resolve_insert_point`, `_resolve_replace_point`, `_resolve_replace_phrase_globally`, `_resolve_append_section_item`, …
  3. `_select_candidate(intent, candidates)`:
     - 1 кандидат → берём
     - ≥2 + embeddings → cosine similarity; если `top_score >= 0.72` и `margin >= 0.08` → **autopick без LLM**
     - Иначе → `_disambiguate_candidate()` — LLM-вызов с `resolver_disambiguation_system.txt`

**Почему layered (regex → embeddings → LLM):**
- Regex-навигация резолвит большинство (точная адресация по `point_number`/`appendix_number`)
- Embeddings — дешёвый способ выбрать из похожих секций
- LLM — последний рубеж, вызывается только при настоящей неоднозначности
""",
    "manual_review": """
**Назначение:** разделить операции на «безопасно применимые» и «требующие ручного просмотра».

**Ключевой код:**
- `redacta/manual_review.py` → `split_operations_for_manual_review(resolved_operations)`
- `SAFE_OPERATION_KINDS = {replace_point, replace_phrase_globally, repeal_point}`
- Заблокировано если:
  - `status != "resolved"`
  - `operation_kind not in SAFE_OPERATION_KINDS`
  - `replace_phrase_globally` без явного global scope

**Почему так:** инсёрты и append'ы в общий список — самые рискованные операции; их легко применить «не туда». Phrase replacement без явного scope может зацепить неожиданное место.
""",
    "edit": """
**Назначение:** детерминированно применить безопасные операции к рабочему документу через OOXML-мутации.

**Ключевой код:**
- `redacta/editor_v2.py` → `EditorV2.edit(base_doc, output_doc, operations)`
  1. `IndexDriftTracker(initial_count=len(paragraphs))` — отслеживает сдвиг индексов после каждой мутации
  2. Сортировка по приоритету:
     - `0`: `replace_appendix_block`
     - `1`: `replace_point`, `replace_phrase_globally`, `append_words_to_point`, `repeal_point`
     - `2`: `insert_point`, `append_section_item`, `insert_list_entry`
  3. Switch по `operation_kind` → `_apply_insert_point`, `_apply_replace_point`, `_apply_replace_phrase_globally`, …

**Ключевые low-level helpers:**
- `replace_paragraph_text_preserving_ooxml(p, text)` — сохраняет `w:rPr` (форматирование)
- `clone_paragraph_after(anchor, text)` — `deepcopy(anchor._p)` + `addnext`
- `_replace_phrase_variants()` — замена с учётом `«»` ↔ `"`, `ё↔е`, dash-вариантов

**Почему здесь нет LLM:** применение правки — детерминированная операция. Если бы LLM управлял редактированием, мы получали бы непредсказуемые мутации и риск утери форматирования.
""",
    "markers": """
**Назначение:** вставить маркеры редакций в стиле «Консультант-плюс» после применённых операций.

**Ключевой код:**
- `redacta/revision_markers.py` → `RevisionMarkerInserter.insert_markers(document_path, operations)`
- `ConsultantMarkerFormatter.format_marker(operation)` — mapping `operation_kind` → строка маркера:
  - `insert_point` → `(п. N введен Приказом ...)`
  - `replace_point` → `(п. N в ред. Приказа ...)`
  - `replace_appendix_block` → `(Приложение N в ред. Приказа ...)`
  - catch-all → `(в ред. Приказа ...)`
- `_introduced_label` — склоняет вид документа в творительный падеж (`Приказ → Приказом`)
- Вставка: сортирует по индексу в обратном порядке, дедуплицирует, проверяет «маркер уже стоит»

**Почему regex, а не LLM:** формат маркера фиксирован; структурные границы ищутся regex'ом (`^\\d+\\.\\s+`, `^[а-я]\\)\\s+`).
""",
    "reanalyze": """
**Назначение:** перестроить структурную карту уже отредактированного документа и обновить служебные таблицы.

**Ключевой код:**
- `BaseAnalyzer.analyze(current_doc)` — повторный анализ финального документа
- `build_service_table_specs(final_base_analysis, amendment_analyses)` — перестроение specs
- `Document.save(current_doc)` — сохранение

**Зачем нужно:** после insert/delete/replace индексы параграфов и структура могли измениться; финальные service tables должны учитывать актуальное состояние.
""",
    "checklist": """
**Назначение:** собрать структурированный `ValidationChecklist` для передачи валидатору.

**Ключевой код:**
- `redacta/validation_checklist_builder.py` → `ValidationChecklistBuilder.build(...)`
- Собирает:
  1. Все runtime-чеки из `PipelineChecklist`
  2. На каждый `ServiceTableSpec` → `service_table_present`
  3. На каждый `ChangeIntent` → детальный check
  4. Глобальный `resolution_completeness` (missing/extra/non-resolved)
  5. На каждую `ResolvedOperation` → `resolved_<id>`

**Почему отдельный builder:** агрегация чеков от разных стадий в единый контракт для валидатора. Ни LLM, ни regex — чистая агрегация.
""",
    "validation": """
**Назначение:** финальная двухуровневая проверка корректности итогового документа.

**Ключевой код:**
- `redacta/validator.py` → `StrictJudgeValidator.validate(...)`

**Уровни проверки:**
1. **Structural checks (regex/substring)**:
   - `_validate_skeleton_tables()` — проверяет, что в таблицах встречаются ВСЕ `expected_documents`
2. **LLM Judge**:
   - 1 вызов LLM с `judge_system.txt` + `judge_user_template.txt`
   - Возвращает `{is_valid, summary, failures}`
3. **Hard fail — resolution mismatch**:
   - Если `total_intents > 0`, но `resolved != total_intents` или `ambiguous > 0`
4. **Deterministic intent checks (regex + substring)**:
   - `append_words_to_point` → `appended_words` присутствует в output
   - `replace_point` → `new_text` присутствует, `^point_number\\.\\s+` ровно 1 раз
   - `repeal_point` → `Утратил силу` присутствует
5. **Phrase coverage**:
   - `_check_phrase_replacement_coverage()` — парсит `operation_statuses` regex'ом на `0 occurrences`
6. **Recovery logic**:
   - Если LLM сказал `false`, но deterministic checks ОК и нет hard-fail → `judge_ok` восстанавливается

**Почему двухуровневая:**
- **Структурная** — быстрая, точная, не требует LLM
- **LLM-судья** — ловит семантические ошибки («правильный текст, но не туда»)
- **Deterministic checks** — safety net от false-negative LLM
""",
}


# ---------------------------------------------------------------------------
# Sidebar
# ---------------------------------------------------------------------------
_ensure_state()

with st.sidebar:
    st.markdown("## ⚙️ Configuration")
    with st.expander("LLM Settings", expanded=False):
        st.session_state["base_url"] = st.text_input("Base URL", value="http://127.0.0.1:11434/v1")
        st.session_state["api_key"] = st.text_input("API Key", value="ollama", type="password")
        st.session_state["model_name"] = st.text_input(
            "Model", value="hf.co/ai-sage/GigaChat3.1-10B-A1.8B-GGUF:Q4_K_M"
        )
        st.session_state["temperature"] = st.slider("Temperature", 0.0, 1.0, 0.0, 0.1)
        st.session_state["timeout"] = st.number_input("Timeout (sec)", 30, 3600, 900)
        st.session_state["semantic_ranking"] = st.toggle("Semantic Ranking", value=False)

    st.markdown("---")
    st.markdown("### 📤 Upload Documents")
    base_file = st.file_uploader("Base Document (.docx)", type=["docx"], key="base_upload")
    amendment_file = st.file_uploader("Amendment Document (.docx)", type=["docx"], key="amendment_upload")

    if base_file and amendment_file:
        st.session_state.files_ready = True
        st.session_state.base_file = base_file
        st.session_state.amendment_file = amendment_file
        # Create a persistent workspace directory in temp for this session
        if "workspace_dir" not in st.session_state or not Path(st.session_state.workspace_dir).exists():
            st.session_state.workspace_dir = tempfile.mkdtemp(prefix="redacta_web_")

    st.markdown("---")
    if st.button("🔄 Reset Pipeline", use_container_width=True):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()

    st.markdown(
        '<div style="font-size:0.8rem;color:#9ca3af;">RedActa Step-by-Step Explorer</div>',
        unsafe_allow_html=True,
    )

# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
st.markdown("# 🔬 RedActa Step-by-Step Explorer")
st.caption("Run each pipeline stage individually and inspect intermediate results.")

if not st.session_state.files_ready:
    st.info("👈 Upload both documents in the sidebar to begin.")
    st.stop()

# Header with file names
st.markdown(f"**Base:** `{st.session_state.base_file.name}`  |  **Amendment:** `{st.session_state.amendment_file.name}`")

render_pipeline_map()

# Determine next actionable step
def _next_step() -> str | None:
    for key, _, _ in PIPELINE_STEPS:
        status = st.session_state.step_status.get(key, "pending")
        if status in ("pending", "error"):
            return key
    return None


next_step = _next_step()

# Active step card
st.markdown("---")
if next_step:
    step_title = _step_display_name(next_step)
    st.markdown(f"## ▶️ Next Step: {step_title}")

    # Show previous step results if any
    prev_idx = next((i for i, (k, _, _) in enumerate(PIPELINE_STEPS) if k == next_step), 0) - 1
    if prev_idx >= 0:
        prev_key = PIPELINE_STEPS[prev_idx][0]
        prev_status = st.session_state.step_status.get(prev_key)
        if prev_status in ("success", "warning") and prev_key in STEP_RENDERERS:
            with st.expander(f"📋 Review previous step: {_step_display_name(prev_key)}", expanded=False):
                STEP_RENDERERS[prev_key]()

    # Description of the step logic
    with st.expander(f"📖 Что происходит на шаге «{_step_display_name(next_step)}»?", expanded=True):
        st.markdown(STEP_DESCRIPTIONS.get(next_step, "Описание отсутствует."))

    col_btn, _ = st.columns([1, 4])
    with col_btn:
        run_clicked = st.button(f"🚀 Run {_step_display_name(next_step)}", type="primary", use_container_width=True)

    if run_clicked:
        runner = STEP_RUNNERS[next_step]
        with st.spinner(f"Running {next_step}..."):
            runner()
        st.rerun()

    # Show logs for current step if already run (e.g., error state)
    if st.session_state.step_status.get(next_step) in ("error", "warning", "success"):
        _render_logs(next_step)
else:
    st.success("🎉 All steps completed!")
    # Final review
    if "validation" in st.session_state.step_status:
        with st.expander("Final Validation", expanded=True):
            _render_step_validation()

    # Downloads
    if st.session_state.current_doc and Path(st.session_state.current_doc).exists():
        doc_bytes = Path(st.session_state.current_doc).read_bytes()
        _download_button(
            doc_bytes,
            "redacta_output.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "📄 Download Output Document",
        )
    c_demo, c_full = st.columns(2)
    with c_demo:
        result_json = json.dumps(
            {
                "case_id": st.session_state.case_id,
                "base_doc": str(st.session_state.base_doc),
                "output_doc": str(st.session_state.current_doc),
                "status": st.session_state.step_status.get("validation", "pending"),
                "validation": st.session_state.validation.to_dict() if st.session_state.validation else None,
            },
            ensure_ascii=False,
            indent=2,
        )
        _download_button(result_json, "redacta_result_demo.json", "application/json", "📋 Download Demo JSON", key="dl_demo_final")
    with c_full:
        full_result = _build_full_result()
        _download_json(full_result, "redacta_result_full.json", "📥 Download Full Result JSON", key="dl_full_final")

# Global error display
if st.session_state.pipeline_error:
    st.error(f"Pipeline error: {st.session_state.pipeline_error}")
    with st.expander("Traceback"):
        st.code(traceback.format_exc(), language="python")
